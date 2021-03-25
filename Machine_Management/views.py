from django.shortcuts import render, redirect
from django.contrib import messages
from django.core.exceptions import ObjectDoesNotExist
from .models import *
import Machine_Management
import datetime
import django
from .forms import *
from django.db.models import Q
from django.http import HttpResponse
import xlwt
from docx import Document
from docx.shared import Inches
from collections import Counter
from docx.enum.text import WD_ALIGN_PARAGRAPH

# from .filters import MachineFilter
# from django.http import JsonResponse
# from django.views.decorators.csrf import csrf_exempt
# from django.core import serializers
# import ast

# Create your views here.

# GLOBAL var
User_login, UserRole, List_user_Screen, dict_menu_level, User_org_machine_line, List_user_Screen = None, None, [], {}, None, None  # User Login for use all pages
UserLoginDepartment = None


def signin(request):
    # Functions for Sign In to webapp
    # Templates/signin.html
    global User_login, UserRole, dict_menu_level
    User_login, UserRole, dict_menu_level = None, None, None
    if request.method == "POST":
        # Form Sign In
        if 'signin' in request.POST:
            username = request.POST['inputUser']  # Get var('username') from HTML
            password = request.POST['inputPassword']  # Get var('password') from HTML
            try:                                                                                    # Try connect username and passwd on Model
                user = User.objects.get(username=username, password=password)
                now = datetime.datetime.now()  # Call Datetime now
                datenow = datetime.date.today()                                                                         # Call Date now
                user.last_login_date = now                                                                              # Update last_login to now
                user.save()                                                                                             # Save Update
                if user.start_date > datenow:                                                                           # Check StartDate and DateNow
                    messages.error(request, f'ชื่อผู้ใช้นี้ยังไม่สามารถเข้าสู่ระบบได้ สามารถเข้าได้ในวันที่ {user.start_date}')
                    return redirect('signin')
                elif datenow > user.expired_date:                                                   # Check Expired Date if expired link to resetpassword
                    messages.info(request, 'รหัสผ่านหมดอายุแล้ว กรุณาทำการ Reset Password')
                    return redirect('signin')
                if user is not None:                                                                # Check login User   # Set user login
                    User_login = user
                    UserRole = str(User_login.role)
                    return redirect('signIn_department')
            except Machine_Management.models.User.DoesNotExist:  # Message Wrong username or password
                messages.error(request, "username หรือ password ไม่ถูกต้อง")

    return render(request, 'account/signin.html')


def usermanage(request):
    # Functions for User Management
    # Templates/usermanage.html
    global User_login, List_user_Screen  # Call User sign in
    if not Role_Screen.objects.filter(role=UserRole, screen_id='usermanage').exists():
        return redirect('signin')
    if request.method == "POST":
        # Form Edituser (Settings of user)
        if 'Edituser' in request.POST:
            username = request.POST['set_username']                 # Get var('username') from HTML
            update_role = request.POST['select_role']               # Get var('role') from HTML
            update_org = request.POST['select_org']
            now = datetime.datetime.now()                           # Call Datetime now
            user = User.objects.get(username=username)              # Query user
            user.user_active = request.POST.get('set_user_status', False)
            user.update_date = now                                  # Update UpdateDate to now
            user.update_by = str(User_login.username)               # Update UserUpdate of UserSelect
            org = Organization.objects.get(org_id=update_org)
            user.org = org
            role = Role.objects.get(role_id=update_role)            # Get RoleID of UserSelect
            user.role = role                                        # Update Role of UserSelect
            user.save()                                             # Save all Update
            messages.success(request, "แก้ไขและบันทึกรายการสำเร็จ")
        # Form Add User (Add New User)
        elif 'Adduser' in request.POST:
            username = request.POST['add_username']                 # Get var('username') form HTML
            fname = request.POST['add_fname']
            lname = request.POST['add_lname']
            email = request.POST['add_email']
            startdate = request.POST['add_startdate']
            create_role = request.POST['select_role']
            passwd = request.POST['add_password']
            conpasswd = request.POST['add_cfpassword']
            add_org = request.POST['add_select_org']
            department = Department.objects.get(pk=request.POST['select_department'])
            now = datetime.datetime.now()
            org = Organization.objects.get(org_id=add_org)
            role = Role.objects.get(role_id=create_role)
            if passwd == conpasswd:  # Check password and confirm password
                if User.objects.filter(username=username).exists():  # Query username is exists in model(DB)
                    messages.error(request, "มีผู้ใช้ชื่อ Username นี้แล้ว")  # Show Message Username is exists
                elif User.objects.filter(email=email).exists():  # Query email is exists in model(DB)
                    messages.error(request, "มีผู้ใช้ Email นี้แล้ว")  # Show Message Email is exists
                else:
                    user = User.objects.create(
                        username=username,
                        email=email,
                        firstname=fname.capitalize(),
                        lastname=lname.capitalize(),
                        password=passwd,
                        create_by=str(User_login.username),
                        create_date=now,
                        expired_date=now - datetime.timedelta(1),
                        expired_day=90,
                        start_date=startdate,
                        update_by=None,
                        update_date=None,
                        last_login_date=None,
                        role=role,
                        org=org,
                        user_active=True,
                    )
                    user.save()         # Save User
                    user.departments.add(department)
                    messages.success(request, "สร้างรายการสำเร็จ")
                    # return redirect('/usermanage')
            else:
                messages.error(request, "รหัสผ่านไม่ตรงกัน กรุณาตรวจสอบใหม่")  # Show Message Password != ConPassword

        # Button Sign Out
        elif 'signout' in request.POST:
            User_login = None  # Set User_login is None

        # Form DeleteUser (icon delete)
        elif 'deleteuser' in request.POST:
            username = request.POST['deleteuser']                               # get var('username') from HTML
            user = User.objects.get(username=username)                          # Query Username
            user.delete()                                                       # Delete User from Model(DB)
            messages.success(request, "ลบรายการสำเร็จ")

        return redirect('usermanage')

    # return var to HTML
    roles = Role.objects.all()
    users = User.objects.all()
    orgs = Organization.objects.all()
    departments = Department.objects.all()
    context = {'users': users,
               'roles': roles,
               'User_login': User_login,
               'orgs': orgs,
               'departments': departments}
    return render(request, 'account/usermanage.html', context)


def reset_password(request):
    # Form Reset Password
    # Get variables from Input HTML
    if request.method == "POST":
        username = request.POST['inputUser']
        old_password = request.POST['oldPassword']
        new_password = request.POST['newPassword']
        con_new_password = request.POST['conPassword']
        try:  # Test connect User in modals(DB)
            user = User.objects.get(username=username, password=old_password)
            if old_password != new_password:
                if new_password == con_new_password:  # Check new_pass and con_pass
                    user.password = new_password
                    now = datetime.date.today()
                    user.expired_date = now + datetime.timedelta(90)
                    user.save()
                    messages.success(request, "สร้างรายการสำเร็จ")
                    return redirect('signin')
                else:  # NewPassword != ConPassword
                    messages.error(request, 'รหัสผ่านใหม่และรหัสผ่านยืนยันไม่ตรงกัน')
            else:  # OldPassword != NewPassword
                messages.error(request, 'รหัสผ่านเก่าต้องไม่ตรงกับรหัสผ่านใหม่')
        except Machine_Management.models.User.DoesNotExist:  # Failed Connect User in model(DB)
            messages.error(request, 'ชื่อผู้ใช้และรหัสผ่านเก่าไม่ถูกต้อง')
    return render(request, 'account/resetpassword.html')


def rolemanage(request):
    global User_login  # Call User sign in
    if not Role_Screen.objects.filter(role=UserRole, screen_id='rolemanage').exists():
        return redirect('signin')
    if request.method == "POST":
        if 'Editrole' in request.POST:
            role_id = request.POST['set_roleid']  # Get var('role id') from HTML
            role_name = request.POST['set_rolename']  # Get var('role name') from HTML
            role = Role.objects.get(role_id=role_id)
            role.role_name = role_name
            role.save()
            messages.success(request, "แก้ไขและบันทึกรายการสำเร็จ")
        elif 'Addrole' in request.POST:
            role_id = request.POST['add_roleid']  # Get var('role id') from HTML
            role_name = request.POST['add_rolename']  # Get var('role name') from HTML
            try:
                role = Role.objects.create(role_id=role_id, role_name=role_name)
                role.save()
                messages.success(request, "สร้างรายการสำเร็จ")
            except django.db.utils.IntegrityError:
                messages.error(request, "มีชื่อ Role ID นี้แล้ว กรุณาตั้งใหม่")
        elif 'deleterole' in request.POST:
            role_id = request.POST['deleterole']  # Get var('role id') from HTML
            role = Role.objects.get(role_id=role_id)
            role.delete()
            messages.success(request, "ลบรายการสำเร็จ")
        elif 'signout' in request.POST:
            User_login = None  # Set User_login is None

        return redirect('rolemanage')

    roles = Role.objects.all()
    context = {'roles': roles,
               'User_login': User_login}
    return render(request, 'account/rolemanage.html', context)


def screenmanage(request):
    global User_login
    if not Role_Screen.objects.filter(role=UserRole, screen_id='screenmanage').exists():
        return redirect('signin')
    if request.method == "POST":
        if 'Addscreen' in request.POST:
            if not Screen.objects.filter(screen_id=request.POST['screen_id']).exists():
                screen_id = request.POST['screen_id']
                screen_name = request.POST['screen_name']
                file_py = request.POST['file_py']
                file_html = request.POST['file_html']
                screen = Screen.objects.create(
                    screen_id=screen_id,
                    screen_name=screen_name,
                    file_py=file_py,
                    file_html=file_html
                )
                screen.save()
                messages.success(request, "สร้างรายการสำเร็จ")
            else:
                messages.info(request, "มีชื่อ Screen ID นี้แล้ว กรุณาตั้งใหม่")
        elif 'deletescreen' in request.POST:
            del_screen_id = request.POST['deletescreen']
            screen = Screen.objects.get(screen_id=del_screen_id)
            screen.delete()
            messages.success(request, "ลบรายการสำเร็จ")
        elif 'Editscreen' in request.POST:
            set_screen_id = request.POST['set_screenid']
            set_screen_name = request.POST['set_screenname']
            set_file_py = request.POST['set_filepy']
            set_file_html = request.POST['set_filehtml']
            screen = Screen.objects.get(screen_id=request.POST['Editscreen'])
            screen.screen_id = set_screen_id
            screen.screen_name = set_screen_name
            screen.file_py = set_file_py
            screen.file_html = set_file_html
            screen.save()
            messages.success(request, "แก้ไขและบันทึกรายการสำเร็จ")

        return redirect('screenmanage')

    screens = Screen.objects.all()
    context = {'User_logined': User_login,
               'screens': screens}
    return render(request, 'account/screenmanage.html', context)


def role_screen(request):
    global User_login
    if not Role_Screen.objects.filter(role=UserRole, screen_id='role_screen').exists():
        return redirect('signin')
    if request.method == "POST":
        if 'delete_rs' in request.POST:
            rs_id = request.POST['delete_rs']
            role_screen = Role_Screen.objects.get(id=rs_id)
            role_screen.delete()
            messages.success(request, "ลบรายการสำเร็จ")
        elif 'Edit_rs' in request.POST:
            rs = Role_Screen.objects.get(id=request.POST['Edit_rs'])
            if rs.role_id == request.POST['set_rs_role'] and rs.screen_id == request.POST['set_rs_screen']:
                rs_id = request.POST['Edit_rs']
                rs_role_id = request.POST['set_rs_role']
                rs_screen_id = request.POST['set_rs_screen']
                rs_insert = request.POST.get('set_rs_insert', "N")
                rs_update = request.POST.get('set_rs_update', "N")
                rs_delete = request.POST.get('set_rs_delete', "N")
                role_screen = Role_Screen.objects.get(id=rs_id)
                role_screen.role_id = rs_role_id
                role_screen.screen_id = rs_screen_id
                role_screen.permission_insert = rs_insert
                role_screen.permission_update = rs_update
                role_screen.permission_delete = rs_delete
                role_screen.save()
                messages.success(request, "แก้ไขและบันทึกรายการสำเร็จ")
            elif not Role_Screen.objects.filter(role_id=request.POST['set_rs_role'],
                                                screen_id=request.POST['set_rs_screen']).exists():
                rs_id = request.POST['Edit_rs']
                rs_role_id = request.POST['set_rs_role']
                rs_screen_id = request.POST['set_rs_screen']
                rs_insert = request.POST.get('set_rs_insert', "N")
                rs_update = request.POST.get('set_rs_update', "N")
                rs_delete = request.POST.get('set_rs_delete', "N")
                role_screen = Role_Screen.objects.get(id=rs_id)
                role_screen.role_id = rs_role_id
                role_screen.screen_id = rs_screen_id
                role_screen.permission_insert = rs_insert
                role_screen.permission_update = rs_update
                role_screen.permission_delete = rs_delete
                role_screen.save()
                messages.success(request, "แก้ไขและบันทึกรายการสำเร็จ")
            else:
                messages.error(request, 'Role และ Screen นี้มีแล้ว!! ไม่สามารถสร้างซ้ำได้')
        elif 'Addrolescreen' in request.POST:
            if not Role_Screen.objects.filter(role_id=request.POST['add_rs_role'],
                                              screen_id=request.POST['add_rs_screen']).exists():
                rs_role_id = request.POST['add_rs_role']
                rs_screen_id = request.POST['add_rs_screen']
                rs_insert = request.POST.get('add_rs_insert', "N")
                rs_update = request.POST.get('add_rs_update', "N")
                rs_delete = request.POST.get('add_rs_delete', "N")
                role_screen = Role_Screen.objects.create(
                    role_id=rs_role_id,
                    screen_id=rs_screen_id,
                    permission_insert=rs_insert,
                    permission_update=rs_update,
                    permission_delete=rs_delete
                )
                role_screen.save()
                messages.success(request, "สร้างรายการสำเร็จ")
            else:
                messages.error(request, 'Role และ Screen นี้มีแล้ว!! ไม่สามารถสร้างซ้ำได้')

        return redirect('role_screen_manage')

    list_role_screen = Role_Screen.objects.all()
    roles = Role.objects.all()
    screens = Screen.objects.all()
    context = {'User_login': User_login,
               'list_role_screen': list_role_screen,
               'roles': roles,
               'screens': screens}
    return render(request, 'account/role_screen_manage.html', context)


def home(request):
    global User_login, dict_menu_level, List_user_Screen
    if request.method == "POST":
        if 'signout' in request.POST:
            User_login = None
    try:
        user_role = Role.objects.get(role_id=User_login.role)
    except AttributeError:
        return redirect("signin")

    user_dep = User_and_department.objects.get(user=User_login, department=UserLoginDepartment)
    if dict_menu_level is None or dict_menu_level == {}:
        List_user_Screen = user_role.members.all()
        list_user_menu_lv0 = Menu.objects.filter(level=0).order_by('index')
        list_user_menu_lv1 = Menu.objects.filter(level=1).order_by('index')
        list_menu_role = []
        dict_menu_level = {}
        for menu_role in List_user_Screen:
            try:
                list_menu_role.append(Menu.objects.get(screen=menu_role))
            except ObjectDoesNotExist:
                pass
        for root in list_user_menu_lv0:
            if root in list_menu_role:
                dict_menu_level[root] = []
        for child in list_user_menu_lv1:
            if child.pk == 'repair_notice' and not user_dep.is_inform and not user_dep.is_close:
                continue
            elif child.pk == 'repair_inspect' and not user_dep.is_inspect:
                continue
            elif child.pk == 'repair_approve' and not user_dep.is_approve:
                continue
            elif child.pk == 'maintenance_receive' and not user_dep.is_receive:
                continue
            elif child.pk == 'maintenance_assign' and not user_dep.is_assign:
                continue
            elif child.pk == 'maintenance_report' and not user_dep.is_report and not user_dep.is_verify:
                continue
            if child in list_menu_role:
                root = Menu.objects.get(menu_id=child.parent_menu)
                dict_menu_level[root].append(child)

    if request.method == "POST":
        if "assign_repair_job" in request.POST:
            engineer_pk = request.POST.get('engineer_username', False)
            repair_notice_model = Repair_notice.objects.get(pk=request.POST['assign_repair_job'])
            repair_notice_model.repair_status = 'อยู่ในระหว่างการทำงาน'

            list_check = request.POST.getlist('select_sp_name1')
            if len(list_check) != len(set(list_check)):
                messages.error(request, "ทำรายการไม่สำเร็จ เนื่องจากเลือกชิ้นส่วนอะไหล่ซ้ำ")
                return redirect('home')
            elif "0" in list_check:
                messages.error(request, "ทำรายการไม่สำเร็จ เนื่องจากไม่ได้เลือกชิ้นส่วนอะไหล่")
                return redirect('home')
            for sp_id_test in list_check:
                try:
                    mch_sp = Machine_sparepart.objects.get(machine=repair_notice_model.machine, spare_part_id=sp_id_test)
                except ObjectDoesNotExist:
                    messages.error(request, "ทำรายการไม่สำเร็จ เนื่องจากเครื่องจักรยังไม่ได้เชื่อมต่อกับอะไหล่")
                    return redirect('home')
                if mch_sp.gen_mtnchng_date is not None:
                    job_auto = Maintenance_job.objects.exclude(job_status__in=["ปิดงาน", "งานเสร็จสิ้น"]).get(job_mch_sp=mch_sp, job_gen_date=mch_sp.gen_mtnchng_date)
                    job_auto.job_status = "ปิดงาน"
                    job_auto.job_remark = "เนื่องจากมีการสร้างงานนี้ในใบแจ้งซ่อม"
                    job_auto.is_report = True
                    job_auto.is_approve = True
                    job_auto.save()
                elif mch_sp.gen_mtnchk_date is not None:
                    job_auto = Maintenance_job.objects.exclude(job_status__in=["ปิดงาน", "งานเสร็จสิ้น"]).get(job_mch_sp=mch_sp, job_gen_date=mch_sp.gen_mtnchk_date)
                    job_auto.job_status = "ปิดงาน"
                    job_auto.job_remark = "เนื่องจากมีการสร้างงานนี้ในใบแจ้งซ่อม"
                    job_auto.is_report = True
                    job_auto.is_approve = True
                    job_auto.save()
                mch_sp.gen_mtnchng_date = datetime.date.today()
                mch_sp.save()
                mtn_repair_job = Maintenance_job.objects.create(job_no=genJobNumber(),
                                                                job_assign_user=User_login,
                                                                job_response_user_id=engineer_pk,
                                                                job_assign_date=datetime.date.today(),
                                                                job_status="รอการดำเนินงาน",
                                                                job_mtn_type="repair",
                                                                job_gen_user=User_login,
                                                                job_mch_sp=mch_sp)
                repair_notice_model.maintenance_jobs.add(mtn_repair_job)
                mtn_repair_job.save()
            repair_notice_model.save()
            messages.success(request, "บันทึกรายการสำเร็จ")
            return redirect('home')
        elif "assign_mtn_job" in request.POST:
            job = Maintenance_job.objects.get(pk=request.POST['assign_mtn_job'])
            if job.job_status in ['รอการอนุมัติงาน', 'งานเสร็จสิ้น', 'ปิดงาน']:
                messages.error(request, 'ไม่สามารถมอบหมายงานได้ เนื่องจากงานอยู่ในสถานะรอการอนุมัติงานแล้วหรืองานเสร็จสิ้น')
                return redirect('home')
            job.job_assign_user = User_login
            job.job_response_user = User.objects.get(pk=request.POST['engineer_username'])
            job.job_assign_date = datetime.date.today()
            job.job_status = "รอการดำเนินงาน"
            job.save()
            return redirect('home')
        elif 'report_submit' in request.POST:
            mtn_report = Maintenance_job.objects.get(pk=request.POST['report_submit'])
            mch_sp = Machine_sparepart.objects.get(pk=mtn_report.job_mch_sp_id)
            mtn_report.job_mtn_type = request.POST['mtn_type']
            mtn_report.job_result_type = request.POST['mtn_result']
            mtn_report.job_result_description = request.POST['mtn_result_description'] if request.POST['mtn_result_description'] != "" else None
            mtn_report.job_fix_plan_hour = request.POST['chang_life_hour_sp'] if request.POST['chang_life_hour_sp'] != "" else request.POST['chang_life_hour_pv']
            mtn_report.job_mch_hour = request.POST['machine_hour_sp'] if request.POST['machine_hour_sp'] != "" else request.POST['machine_hour_pv']
            mtn_report.job_plan_hour = request.POST['check_life_hour_sp'] if request.POST['check_life_hour_sp'] != "" else request.POST['check_life_hour_pv']
            mtn_report.job_report_date = datetime.datetime.today()
            mtn_report.problem_cause = request.POST['problem_cause']
            mtn_report.corrective_action = request.POST['corrective_action']
            mtn_report.after_repair = request.POST['after_repair']
            mtn_report.job_status = "รอการอนุมัติงาน"
            mtn_report.is_approve = False
            mtn_report.is_report = True
            mtn_report.equipment_code1 = request.POST['equipment_code1'] if request.POST['equipment_code1'] != "" else None
            mtn_report.equipment_code2 = request.POST['equipment_code2'] if request.POST['equipment_code2'] != "" else None
            mtn_report.equipment_code3 = request.POST['equipment_code3'] if request.POST['equipment_code3'] != "" else None
            mtn_report.equipment_detail1 = request.POST['equipment_detail1'] if request.POST['equipment_detail1'] != "" else None
            mtn_report.equipment_detail2 = request.POST['equipment_detail2'] if request.POST['equipment_detail2'] != "" else None
            mtn_report.equipment_detail3 = request.POST['equipment_detail3'] if request.POST['equipment_detail3'] != "" else None
            mtn_report.equipment_quantity1 = request.POST['equipment_quantity1'] if request.POST['equipment_quantity1'] != "" else None
            mtn_report.equipment_quantity2 = request.POST['equipment_quantity2'] if request.POST['equipment_quantity2'] != "" else None
            mtn_report.equipment_quantity3 = request.POST['equipment_quantity3'] if request.POST['equipment_quantity3'] != "" else None
            mtn_report.equipment_note1 = request.POST['equipment_note1'] if request.POST['equipment_note1'] != "" else None
            mtn_report.equipment_note2 = request.POST['equipment_note2'] if request.POST['equipment_note2'] != "" else None
            mtn_report.equipment_note3 = request.POST['equipment_note3'] if request.POST['equipment_note3'] != "" else None
            mtn_report.job_remark = request.POST['job_remark'] if request.POST['job_remark'] != "" else None
            mtn_report.estimate_cost = request.POST['estimate_cost'] if request.POST['estimate_cost'] != "" else None
            try:
                mtn_report.save()
                mch_sp.save()
            except ValueError:
                messages.error(request, 'กรุณากรอกชั่วโมงการเปลี่ยนและชั่วโมงการตรวจสอบของอะไหล่')
            return redirect('home')
        elif "approve_job" in request.POST:
            if request.POST.get('is_approve', False):
                mtn_report = Maintenance_job.objects.get(pk=request.POST['approve_job'])
                mch_sp = Machine_sparepart.objects.get(pk=mtn_report.job_mch_sp_id)
                mtn_report.is_approve = True
                mtn_report.job_approve_date = datetime.datetime.today()
                mtn_report.job_status = "งานเสร็จสิ้น" if request.POST.get('is_approve', False) else "รอการอนุมัติงาน"
                mch_sp.gen_mtnchng_date = None
                mch_sp.gen_mtnchk_date = None
                machine_model = Machine.objects.get(pk=mch_sp.machine.pk)
                machine_model.machine_hour_last_update = mch_sp.machine.machine_hour
                machine_model.machine_hour = mtn_report.job_mch_hour

                if mtn_report.job_mtn_type in ["change", "repair"]:
                    mch_sp.last_mtnchk_hour = mtn_report.job_mch_hour
                    mch_sp.last_mtnchng_hour = mtn_report.job_mch_hour
                    mch_sp.last_mtnchng_job_id = mtn_report.id
                elif mtn_report.job_mtn_type == "checking":
                    mch_sp.last_mtnchk_hour = mtn_report.job_mch_hour
                    mch_sp.last_mtnchk_job_id = mtn_report.id
                mch_sp.mtnchng_life_hour = mtn_report.job_fix_plan_hour
                mch_sp.mtnchk_life_hour = mtn_report.job_plan_hour
                if mch_sp.mtnchk_life_hour:
                    mch_sp.next_mtnchk_hour = int(mch_sp.last_mtnchk_hour) + int(mch_sp.mtnchk_life_hour)
                if mch_sp.mtnchng_life_hour:
                    mch_sp.next_mtnchng_hour = int(mch_sp.last_mtnchng_hour) + int(mch_sp.mtnchng_life_hour)

                machine_model.save()
                mch_sp.save()
                mtn_report.save()

            else:
                mtn_report = Maintenance_job.objects.get(pk=request.POST['approve_job'])
                mtn_report.job_status = "ไม่ผ่านการอนุมัติ"
                mtn_report.is_report = False
                mtn_report.save()
            return redirect('home')

    user_org = User_login.org.org_line.all()
    engineers = User.objects.filter(role__role_id__contains="engineer")
    list_repair_notice = Repair_notice.objects.filter(repair_status='รอการตรวจสอบอะไหล่')
    maintenance_job_gen = Maintenance_job.objects.filter(~Q(job_status='งานเสร็จสิ้น'))
    spare_part_group_all = Spare_part_group.objects.all()

    today = datetime.date.today()
    today_request_tasks = Repair_notice.objects.filter(repair_gen_date=today).exclude(repair_status='ปิดใบแจ้งซ่อม')
    today_maintenance_tasks = Maintenance_job.objects.filter(~Q(job_status='งานเสร็จสิ้น'), job_gen_date=today)
    repair_inform_incomplete = Repair_notice.objects.filter(repairer_user=User_login).exclude(repair_status='ปิดใบแจ้งซ่อม')
    mtn_report_incomplete = Maintenance_job.objects.filter(Q(job_response_user=User_login), Q(job_status='รอการดำเนินงาน') | Q(job_status='รอการอนุมัติงาน'))

    context = {'User_login': User_login, 'dict_menu_level': dict_menu_level.items(),
               'line_of_user': user_org, 'user_dep': user_dep,
               'repair_receive': list_repair_notice, 'maintenance_job_gen': maintenance_job_gen,
               'today_request_tasks': today_request_tasks, 'today_maintenance_tasks': today_maintenance_tasks,
               'repair_inform_incomplete': repair_inform_incomplete, 'mtn_report_incomplete': mtn_report_incomplete,
               'engineers': engineers, 'spare_part_group_all': spare_part_group_all}

    return render(request, 'home/home.html', context)


def menumanage(request):
    global User_login
    if not Role_Screen.objects.filter(role=UserRole, screen_id='menumanage').exists():
        return redirect('signin')
    if request.method == 'POST':
        if 'Addmenu' in request.POST:
            if not Menu.objects.filter(menu_id=request.POST['add_menu_id']).exists():
                add_menu_id = request.POST['add_menu_id']
                add_menu_name = request.POST['add_menu_name']
                add_menu_level = request.POST['add_menu_level']
                select_screen = request.POST['select_screen']
                select_parent = request.POST['select_parent'] if request.POST.get('select_parent', False) else None
                add_menu_index = request.POST['add_menu_index']
                add_menu_path = request.POST['add_menu_path']
                screen = Screen.objects.get(screen_id=select_screen)
                menu = Menu.objects.create(
                    menu_id=add_menu_id,
                    name=add_menu_name,
                    level=add_menu_level,
                    screen=screen,
                    parent_menu=select_parent,
                    index=add_menu_index,
                    path_url=add_menu_path
                )
                menu.save()
                messages.success(request, "สร้างรายการสำเร็จ")
            else:
                messages.error(request, "มีการใช้ Menu ID นี้แล้ว กรุณาตั้งชื่อใหม่")
        elif 'Editmenu' in request.POST:
            old_menu_id = request.POST['Editmenu']
            set_menu_id = request.POST['set_menu_id']
            set_menu_name = request.POST['set_menu_name']
            set_menu_level = request.POST['set_menu_level']
            select_screen = request.POST['select_screen']
            select_parent = request.POST['select_parent']
            set_menu_index = request.POST['set_menu_index']
            set_menu_path = request.POST['set_menu_path']
            screen = Screen.objects.get(screen_id=select_screen)
            if old_menu_id == set_menu_id:
                set_menu = Menu.objects.get(menu_id=set_menu_id)
                set_menu.name = set_menu_name
                set_menu.level = set_menu_level
                set_menu.screen = screen
                set_menu.parent_menu = select_parent
                set_menu.index = set_menu_index
                set_menu.path_url = set_menu_path
                set_menu.save()
                messages.success(request, "แก้ไขและบันทึกรายการสำเร็จ")
            else:
                old_menu = Menu.objects.get(menu_id=old_menu_id)
                old_menu.delete()
                set_menu = Menu.objects.create(
                    menu_id=set_menu_id,
                    name=set_menu_name,
                    level=set_menu_level,
                    screen=screen,
                    parent_menu=select_parent,
                    index=set_menu_index,
                    path_url=set_menu_path
                )
                set_menu.save()
                messages.success(request, "แก้ไขและบันทึกรายการสำเร็จ")
        elif 'deletemenu' in request.POST:
            del_menu_id = request.POST['deletemenu']
            menu_del = Menu.objects.get(menu_id=del_menu_id)
            menu_del.delete()
            messages.success(request, "ลบรายการสำเร็จ")

        return redirect('menumanage')

    list_menu = Menu.objects.order_by('level', 'index')
    list_screen = Screen.objects.all()
    screen_of_menu = []
    for menu in list_menu:
        screen_of_menu.append(menu.screen_id)
    context = {
        'User_login': User_login,
        'list_menu': list_menu,
        'list_screen': list_screen,
        'screen_of_menu': screen_of_menu
    }
    return render(request, 'account/menumanage.html', context)


def organizemanage(request):
    global User_login
    if not Role_Screen.objects.filter(role=UserRole, screen_id='organize_manage').exists():
        return redirect('signin')
    if request.method == 'POST':
        if 'Addorg' in request.POST:
            if not Organization.objects.filter(org_code=request.POST['add_org_code']).exists():
                add_org_code = request.POST['add_org_code']
                add_org_name = request.POST['add_org_name']
                organize = Organization.objects.create(
                    org_code=add_org_code,
                    org_name=add_org_name
                )
                organize.save()
                messages.success(request, "สร้างรายการสำเร็จ")
            else:
                messages.error(request, 'มีชื่อ Organize Code นี้แล้ว ไม่สามารถเพิ่มได้ กรุณาทำรายการใหม่')
        elif 'delete_org' in request.POST:
            organize = Organization.objects.get(org_id=request.POST['delete_org'])
            organize.delete()
            messages.success(request, "ลบรายการสำเร็จ")
        elif 'Editorg' in request.POST:
            organize = Organization.objects.get(org_id=request.POST['set_org_id'])
            organize.org_code = request.POST['set_org_code']
            organize.org_name = request.POST['set_org_name']
            organize.save()
            messages.success(request, "แก้ไขและบันทึกรายการสำเร็จ")

        return redirect('organization')

    orgs = Organization.objects.all()
    context = {
        'orgs': orgs, 'User_login': User_login
    }
    return render(request, 'organization/organizemanage.html', context)


def production_line(request):
    global User_login, UserRole
    if not Role_Screen.objects.filter(role=UserRole, screen_id='production_line').exists():
        return redirect('signin')
    if request.method == "POST":
        if 'Addprodline' in request.POST:
            try:
                if not Production_line.objects.filter(production_line=request.POST['add_prodline'],
                                                      location_site=Site.objects.get(
                                                          id=request.POST['add_select_site']),
                                                      location_building=Building.objects.get(
                                                          id=request.POST['add_select_building']),
                                                      location_floor=Floor.objects.get(
                                                          id=request.POST['add_select_floor'])
                                                      ).exists():
                    pline = Production_line.objects.create(
                        production_line=request.POST['add_prodline'],
                        location_site=Site.objects.get(id=request.POST['add_select_site']),
                        location_building=Building.objects.get(id=request.POST['add_select_building']),
                        location_floor=Floor.objects.get(id=request.POST['add_select_floor'])
                    )
                    pline.save()
                    messages.success(request, "สร้างรายการสำเร็จ")
                else:
                    messages.error(request, "มี Production Line นี้แล้วอยู่ในระบบ")
            except Machine_Management.models.Floor.DoesNotExist:
                messages.error(request, "คุณกรอกข้อมูลบางส่วนไม่สมบูรณ์ กรุณากรอกข้อมูลให้สมบูรณ์")

        elif 'Edit_prod_line' in request.POST:
            if not Production_line.objects.filter(production_line=request.POST['set_production_line'],
                                                  location_site__site=request.POST['select_site'],
                                                  location_building__building=request.POST['select_building'],
                                                  location_floor__floor=request.POST['select_floor']).exists():
                pline = Production_line.objects.get(pid=request.POST['set_prodline_id'])
                pline.production_line = request.POST['set_production_line']
                pline.save()
            else:
                messages.error(request, "มี Production Line นี้แล้วอยู่ในระบบ")
        elif 'delete_line' in request.POST:
            pline = Production_line.objects.get(pid=request.POST['delete_line'])
            pline.delete()
            messages.success(request, "ลบรายการสำเร็จ")

        return redirect('productionline')

    lines = Production_line.objects.all()
    sites = Site.objects.all()
    buildings = Building.objects.all()
    floors = Floor.objects.all()
    context = {
        'User_login': User_login, 'lines': lines, 'sites': sites, 'buildings': buildings, 'floors': floors
    }
    return render(request, 'organization/production_line.html', context)


def location(request):
    global User_login, UserRole
    if not Role_Screen.objects.filter(role=UserRole, screen_id='location').exists():
        return redirect('signin')
    if request.method == "POST":
        if 'add_location' in request.POST:
            if Site.objects.filter(site=request.POST['add_site']).exists():
                site = Site.objects.get(site=request.POST['add_site'])
            else:
                site = Site.objects.create(site=request.POST['add_site'])
                site.save()
            if Building.objects.filter(building=request.POST['add_building'], site=site).exists():
                building = Building.objects.get(building=request.POST['add_building'], site=site)
            else:
                building = Building.objects.create(building=request.POST['add_building'], site=site)
                building.save()
            if Floor.objects.filter(floor=request.POST['add_floor'], site=site, building=building).exists():
                messages.error(request, "มี Location นี้แล้ว")
            else:
                floor = Floor.objects.create(floor=request.POST['add_floor'], site=site, building=building)
                floor.save()
                messages.success(request, "สร้างรายการสำเร็จ")

        elif 'Editlocation' in request.POST:
            floor = Floor.objects.get(id=request.POST['set_location_id'])
            if floor.site.site != request.POST['set_site']:
                if Site.objects.filter(site=request.POST['set_site']).exists():
                    set_site = Site.objects.get(site=request.POST['set_site'])
                else:
                    if Floor.objects.filter(site_id=floor.site_id).count() == 1:
                        set_site = Site.objects.get(pk=floor.site_id)
                        set_site.site = request.POST['set_site']
                        set_site.save()
                    else:
                        set_site = Site.objects.create(site=request.POST['set_site'])
                        set_site.save()
                floor.site = set_site
            else:
                set_site = floor.site
            if floor.building.building != request.POST['set_building']:
                if Building.objects.filter(building=request.POST['set_building'], site=set_site).exists():
                    set_building = Building.objects.get(building=request.POST['set_building'])
                else:
                    if Floor.objects.filter(building_id=floor.building_id).count() == 1:
                        set_building = Building.objects.get(pk=floor.building_id)
                        set_building.building = request.POST['set_building']
                        set_building.save()
                    else:
                        set_building = Building.objects.create(building=request.POST['set_building'], site=set_site)
                        set_building.save()
                floor.building = set_building
            else:
                set_building = floor.building
            if floor.floor != request.POST['set_floor']:
                if Floor.objects.filter(floor=request.POST['set_floor'], site=set_site, building=set_building).exists():
                    messages.error(request, "มี Location นี้แล้ว")
                else:
                    floor.floor = request.POST['set_floor']
                    floor.save()
            else:
                floor.save()
            messages.success(request, "แก้ไขและบันทึกรายการสำเร็จ")
        elif 'delete_location' in request.POST:
            locations = Floor.objects.get(pk=request.POST['delete_location'])
            if Floor.objects.filter(site_id=locations.site_id).count() == 1:
                site = Site.objects.get(pk=locations.site_id)
                site.delete()
            if Floor.objects.filter(building_id=locations.building_id).count() == 1:
                building = Building.objects.get(pk=locations.building_id)
                building.delete()
            if Floor.objects.filter(pk=locations.pk).count() == 1:
                locations.delete()
            messages.success(request, "ลบรายการสำเร็จ")

        return redirect('location')

    sites = Site.objects.all()
    buildings = Building.objects.all()
    floors = Floor.objects.all()
    context = {'User_login': User_login, 'sites': sites, 'buildings': buildings, 'floors': floors}
    return render(request, 'organization/location.html', context)


def org_productline(request):
    global User_login, UserRole
    if not Role_Screen.objects.filter(role=UserRole, screen_id='location').exists():
        return redirect('signin')
    if request.method == "POST":
        if "Editorgline" in request.POST:
            org = Organization.objects.get(org_id=request.POST["org_id"])
            line = Production_line.objects.get(pid=request.POST["select_line"])
            org.org_line.add(line)
            org.save()
            messages.success(request, "เพิ่มรายการสำเร็จ")
        elif "delete_org" in request.POST:
            org = Organization.objects.get(org_id=request.POST["delete_org"])
            line = Production_line.objects.get(pid=request.POST["select_del_line"])
            org.org_line.remove(line)
            org.save()
            messages.success(request, "ลบรายการสำเร็จ")

        return redirect('org_prodline')

    org_lines = Organization.objects.all()
    prod_lines = Production_line.objects.all()
    context = {
        'org_lines': org_lines, 'prod_lines': prod_lines, 'User_login': User_login, 'UserRole': UserRole
    }
    return render(request, 'organization/org_prodline.html', context)


def productmanage(request):
    global User_login, UserRole
    if request.method == "POST":
        if 'Addpline' in request.POST:
            if not Product.objects.filter(product_code=request.POST['add_product_code']).exists():
                line = Production_line.objects.get(pk=request.POST['add_select_pline'])
                product = Product.objects.create(
                    product_name=request.POST['add_product_name'],
                    product_code=request.POST['add_product_code'],
                    capacity=request.POST['add_product_capacity'],
                    labour=request.POST['add_product_labour'],
                    line=line
                )
                product.save()
                messages.success(request, "สร้างรายการสำเร็จ")
            else:
                messages.info(request, "มีเลขรหัสผลิตภัณฑ์นี้แล้วนในไลน์การผลิต")
        elif 'Editproduct' in request.POST:
            line = Production_line.objects.get(pk=request.POST['set_select_pline'])
            product = Product.objects.get(pk=request.POST['Editproduct'])
            product.product_name = request.POST['set_product_name']
            product.product_code = request.POST['set_product_code']
            product.capacity = request.POST['set_product_capacity']
            product.labour = request.POST['set_product_labour']
            product.line = line
            product.save()
            messages.success(request, "แก้ไขและบันทึกรายการสำเร็จ")
        elif 'delete_line' in request.POST:
            product = Product.objects.get(pk=request.POST['delete_line'])
            product.delete()
            messages.success(request, "ลบรายการสำเร็จ")

        return redirect('productmanage')

    products = Product.objects.all()
    plines = Production_line.objects.all()
    context = {"User_login": User_login, "UserRole": UserRole, "products": products, 'plines': plines}
    return render(request, 'organization/productmanage.html', context)


def machine_manage(request):
    global User_login, UserRole
    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='machine_management')
    if not role_and_screen.exists():
        return redirect('signin')

    mch_menu = dict_menu_level[Menu.objects.get(pk='machine_manage_menu')]
    mch_main_menu = Menu.objects.get(pk='machine_management')

    user_org = User_login.org.org_line.all()
    machine = Machine.objects.filter(line__in=user_org)
    mch_type_all = Machine_type.objects.all()
    mch_subtype_all = Machine_subtype.objects.all()
    pd_line = Production_line.objects.all()

    filter_mch_line = Machine.objects.order_by('line').values_list('line', flat=True).distinct()
    select_line_export = Production_line.objects.filter(pid__in=filter_mch_line)

    if request.method == "POST":
        if 'Addmachine' in request.POST:
            add_production_line = request.POST['add_production_line']
            add_type = request.POST['add_type']
            add_subtype = request.POST['add_subtype']
            add_serial = request.POST['add_serial']
            add_machine_production_line_code = request.POST['add_mpc']
            add_machine_name = request.POST['add_machine_name']
            add_machine_model = request.POST['add_machine_model']
            add_machine_brand = request.POST['add_machine_brand']
            add_supplier = request.POST['add_supplier_code']
            add_supplier_name = request.POST['add_supplier_name']
            add_supplier_contact = request.POST['add_supplier_contact']
            add_eng_emp_id = request.POST['add_eng_emp_id']
            add_eng_emp_name = request.POST['add_eng_emp_name']
            add_eng_emp_contact = request.POST['add_eng_emp_contact']
            add_pro_emp_id = request.POST['add_pro_emp_id']
            add_pro_emp_name = request.POST['add_pro_emp_name']
            add_pro_emp_contact = request.POST['add_pro_emp_contact']
            add_capacity_per_min = request.POST['add_cpm']
            add_capacity = request.POST['add_capacity']
            add_power = request.POST['add_power'] if request.POST['add_power'] != '' else ""
            add_install_date = request.POST['add_installdate']
            add_start_date = request.POST['add_startdate']
            add_hour = request.POST['add_hour'] if request.POST['add_hour'] != '' else 0
            add_core = request.POST.get('add_mch_core', False)
            add_machine_asset_code = request.POST['add_machine_asset_code'] if request.POST['add_machine_asset_code'] != "" else None
            if request.POST.get('add_mch_core', False):
                if Machine.objects.filter(line_id=add_production_line, machine_core=True).exists():
                    messages.error(request, 'ในไลน์ผลิตนี้มี Machine Core แล้วไม่สามารถทำรายการได้')
                    return redirect('machine_manage')
            if not Machine.objects.filter(serial_id=add_serial,
                                          machine_production_line_code=add_machine_production_line_code).exists():
                add_new_machine = Machine.objects.create(
                    serial_id=add_serial,
                    machine_production_line_code=add_machine_production_line_code,
                    machine_name=add_machine_name,
                    machine_brand=add_machine_brand,
                    machine_model=add_machine_model,
                    machine_supplier_code=add_supplier,
                    machine_supplier_name=add_supplier_name,
                    machine_supplier_contact=add_supplier_contact,
                    machine_eng_emp_id=add_eng_emp_id,
                    machine_eng_emp_name=add_eng_emp_name,
                    machine_eng_emp_contact=add_eng_emp_contact,
                    machine_pro_emp_id=add_pro_emp_id,
                    machine_pro_emp_name=add_pro_emp_name,
                    machine_pro_emp_contact=add_pro_emp_contact,
                    machine_load_capacity=add_capacity_per_min,
                    machine_load_capacity_unit=add_capacity,
                    machine_power_use_kwatt_per_hour=add_power,
                    machine_installed_datetime=add_install_date,
                    machine_start_use_datetime=add_start_date,
                    line_id=add_production_line,
                    mch_type_id=add_type,
                    sub_type_id=add_subtype,
                    create_by=User_login.username,
                    create_date=datetime.date.today(),
                    machine_hour=add_hour,
                    machine_active=True,
                    machine_core=add_core,
                    machine_asset_code=add_machine_asset_code
                )
                add_new_machine.save()
                messages.success(request, 'เพิ่มข้อมูล Machine เรียบร้อยแล้ว')
            else:
                messages.error(request, 'การเพิ่มข้อมูล Machine ล้มเหลว กรุณากด Add New Machine Type ใหม่อีกครั้ง')

        elif 'EditMch' in request.POST:
            edit_mch = Machine.objects.get(machine_id=request.POST['EditMch'])
            edit_mch.machine_production_line_code = request.POST['set_mch_code']
            edit_mch.machine_name = request.POST['set_mch_name']
            # edit_mch.machine_brand = request.POST['set_machine_brand']
            # edit_mch.machine_model = request.POST['set_machine_model']
            # edit_mch.serial_id = request.POST['set_serial']
            edit_mch.machine_supplier_code = request.POST['set_supplier_code']
            edit_mch.machine_supplier_name = request.POST['set_supplier_name']
            edit_mch.machine_supplier_contact = request.POST['set_supplier_contact']
            edit_mch.machine_eng_emp_id = request.POST['set_eng_emp_id']
            edit_mch.machine_eng_emp_name = request.POST['set_eng_emp_name']
            edit_mch.machine_eng_emp_contact = request.POST['set_eng_emp_contact']
            edit_mch.machine_pro_emp_id = request.POST['set_pro_emp_id']
            edit_mch.machine_pro_emp_name = request.POST['set_pro_emp_name']
            edit_mch.machine_pro_emp_contact = request.POST['set_pro_emp_contact']
            edit_mch.machine_load_capacity = request.POST['set_cpm']
            edit_mch.machine_load_capacity_unit = request.POST['set_capacity']
            edit_mch.machine_power_use_kwatt_per_hour = request.POST['set_power']
            edit_mch.machine_installed_datetime = request.POST['set_installdate']
            edit_mch.machine_start_use_datetime = request.POST['set_startdate']
            edit_mch.line_id = request.POST['select_line']
            edit_mch.mch_type_id = request.POST['select_type']
            edit_mch.sub_type_id = request.POST['select_subtype']
            edit_mch.last_update_by = str(User_login.username)
            edit_mch.last_update_date = datetime.date.today()
            edit_mch.machine_active = request.POST.get('set_mch_status', False)
            edit_mch.machine_asset_code = request.POST['set_mch_asset_code'] if request.POST['set_mch_asset_code'] != "" else None
            if request.POST.get('set_mch_core', False):
                if Machine.objects.filter(line_id=edit_mch.line_id, machine_core=True).exists():
                    if edit_mch != Machine.objects.get(line_id=edit_mch.line_id, machine_core=True):
                        messages.error(request, 'ในไลน์ผลิตนี้มี Machine Core แล้วไม่สามารถทำรายการได้')
                        return redirect('machine_manage')
            edit_mch.machine_core = request.POST.get('set_mch_core', False)
            edit_mch.machine_hour = request.POST.get('set_hour', 0)
            edit_mch.machine_minute = request.POST.get('set_minute', 0)
            if edit_mch.machine_document1 != request.FILES.get('set_documentFile1', "") and request.FILES.get('set_documentFile1', False):
                edit_mch.machine_document1.delete()
                edit_mch.machine_document1 = request.FILES['set_documentFile1']
            if edit_mch.machine_document2 != request.FILES.get('set_documentFile2', "") and request.FILES.get('set_documentFile2', False):
                edit_mch.machine_document2.delete()
                edit_mch.machine_document2 = request.FILES['set_documentFile2']
            if edit_mch.machine_document3 != request.FILES.get('set_documentFile3', "") and request.FILES.get('set_documentFile3', False):
                edit_mch.machine_document3.delete()
                edit_mch.machine_document3 = request.FILES['set_documentFile3']
            if edit_mch.machine_document4 != request.FILES.get('set_documentFile4', "") and request.FILES.get('set_documentFile4', False):
                edit_mch.machine_document4.delete()
                edit_mch.machine_document4 = request.FILES['set_documentFile2']
            if edit_mch.machine_document5 != request.FILES.get('set_documentFile5', "") and request.FILES.get('set_documentFile5', False):
                edit_mch.machine_document5.delete()
                edit_mch.machine_document5 = request.FILES['set_documentFile5']
            if edit_mch.machine_image1 != request.FILES.get('set_pictureFile1', "") and request.FILES.get('set_pictureFile1', False):
                edit_mch.machine_image1.delete()
                edit_mch.machine_image1 = request.FILES['set_pictureFile1']
            if edit_mch.machine_image2 != request.FILES.get('set_pictureFile2', "") and request.FILES.get('set_pictureFile2', False):
                edit_mch.machine_image2.delete()
                edit_mch.machine_image2 = request.FILES['set_pictureFile2']
            if edit_mch.machine_image3 != request.FILES.get('set_pictureFile3', "") and request.FILES.get('set_pictureFile3', False):
                edit_mch.machine_image3.delete()
                edit_mch.machine_image3 = request.FILES['set_pictureFile3']
            if edit_mch.machine_image4 != request.FILES.get('set_pictureFile4', "") and request.FILES.get('set_pictureFile4', False):
                edit_mch.machine_image4.delete()
                edit_mch.machine_image4 = request.FILES['set_pictureFile4']
            if edit_mch.machine_image5 != request.FILES.get('set_pictureFile5', "") and request.FILES.get('set_pictureFile5', False):
                edit_mch.machine_image5.delete()
                edit_mch.machine_image5 = request.FILES['set_pictureFile5']
            edit_mch.save()
            messages.success(request, 'แก้ไขข้อมูล Machine สำเร็จ')

        elif 'deletemachine' in request.POST:
            del_machine = request.POST['deletemachine']
            machine_id = Machine.objects.get(machine_id=del_machine)
            machine_id.delete()
            messages.success(request, "ลบรายการสำเร็จ")

        elif 'Export_machine' in request.POST:
            list_line = request.POST.getlist('multi_select_lines')
            list_machine = request.POST.get('Export_machine').split(',')
            if list_line == []:
                list_machine = Machine.objects.filter(line_id__in=user_org)
            elif list_machine == ['']:
                list_machine = Machine.objects.filter(line_id__in=list_line)

            file_report = request.POST['file_type']
            if file_report == 'docx':

                line_id = Machine.objects.filter(machine_id__in=list_machine).order_by('line_id').values('line_id').distinct()
                list_production_line = Production_line.objects.filter(pid__in=line_id).order_by('production_line')
                document = Document()
                for line in list_production_line:
                    document.add_heading(f'Production Line {line.production_line}', 0)

                    document.add_paragraph(
                        f'สถานที่ตั้งโรงงาน {line.location_site} อาคารที่ {line.location_building} ชั้นที่ {line.location_floor}')
                    document.add_paragraph('ผลิตภัณฑ์ที่ผลิต', style='List Bullet')

                    list_product = []
                    for product in Product.objects.filter(line_id=line.pk):
                        try:
                            capacity_core = Machine_capacity.objects.get(machine_id=Machine.objects.get(line_id=line, machine_core=1), product_id=product)
                            list_product.append([str(product.product_name), str(product.product_code), int(capacity_core.fg_capacity)])
                        except Machine_Management.models.Machine.DoesNotExist:
                            capacity_core = "-"
                            list_product.append([str(product.product_name), str(product.product_code), capacity_core])
                        except Machine_Management.models.Machine_capacity.DoesNotExist:
                            capacity_core = "-"
                            list_product.append([str(product.product_name), str(product.product_code), capacity_core])

                    table = document.add_table(rows=1, cols=3)
                    table.style = 'Light List Accent 3'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Product Name'
                    hdr_cells[1].text = 'Product Code'
                    hdr_cells[2].text = 'Product Capacity'
                    for name, code, capacity in list_product:
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(name)
                        row_cells[1].text = str(code)
                        row_cells[2].text = str(capacity)

                    # document.add_heading('เครื่องจักร', level=1)
                    machine_in_line = Machine.objects.filter(line_id=line, machine_id__in=list_machine).order_by('machine_production_line_code')
                    for mch in machine_in_line:

                        document.add_heading(f'ชื่อเครื่องจักร : {mch.machine_name}', level=1)
                        if mch.machine_image1:
                            picture_mch = document.add_picture(mch.machine_image1, width=Inches(2))
                            last_paragraph = document.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # run = document.add_paragraph(f'{mch.machine_name}', style='Intense Quote').add_run()
                        # font = run.font
                        # font.size = Pt(20)
                        document.add_paragraph('ข้อมูลเครื่องจักร', style='List Bullet')

                        records = (
                            ('Machine Name', mch.machine_name),
                            ('Machine Brand', mch.machine_brand),
                            ('Machine Model', mch.machine_model),
                            ('Machine Serial', mch.serial_id),
                            ('Machine Type', mch.mch_type),
                            ('Machine Subtype', mch.sub_type),
                            ('Machine Production Line', mch.line),
                            ('Machine Line Code', mch.machine_production_line_code),
                            ('Machine Asset Code', mch.machine_asset_code),
                            ('Machine Load Capacity', str(mch.machine_load_capacity) + " " + str(mch.machine_load_capacity_unit)),
                            ('Machine Power', str(mch.machine_power_use_kwatt_per_hour)+" KWatt/Hour"),
                            ('Machine Installed Date', mch.machine_installed_datetime),
                            ('Machine Start Date', mch.machine_start_use_datetime),
                            ('Machine Hours', str(mch.machine_hour)),
                            ('Machine Supplier', mch.machine_supplier_code),
                            ('Machine Supplier Name', str(mch.machine_supplier_name)+" (ติดต่อ : "+str(mch.machine_supplier_contact)+" )"),
                            ('Engineer Emp in Charge', str(mch.machine_eng_emp_id)+"  "+str(mch.machine_eng_emp_name)+" (ติดต่อ : "+str(mch.machine_eng_emp_contact)+" )"),
                            ('Production Emp in Charge', str(mch.machine_pro_emp_id)+"  "+str(mch.machine_pro_emp_name)+" (ติดต่อ : "+str(mch.machine_pro_emp_contact)+" )")
                        )

                        table = document.add_table(rows=1, cols=2)
                        table.style = 'Light List Accent 1'
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'Title'
                        hdr_cells[1].text = 'Specification'
                        for title, data in records:
                            row_cells = table.add_row().cells
                            row_cells[0].text = title
                            row_cells[1].text = str(data)

                        mch_capacity = Machine_capacity.objects.filter(machine_id=mch.machine_id)
                        if mch_capacity.exists():
                            document.add_paragraph(' ')
                            document.add_paragraph('ข้อมูลกำลังการผลิต', style='List Bullet')
                            list_table = []
                            for mch_cap in mch_capacity:
                                list_table.append([str(mch_cap.product.product_name), str(mch_cap.product.product_code), str(int(mch_cap.fg_capacity))])
                            table = document.add_table(rows=1, cols=3)
                            table.style = 'Light List Accent 2'
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = 'Product Name'
                            hdr_cells[1].text = 'Product Code'
                            hdr_cells[2].text = 'FG Capacity/Hour'
                            for name, code, capacity in list_table:
                                row_cells = table.add_row().cells
                                row_cells[0].text = name
                                row_cells[1].text = code
                                row_cells[2].text = capacity

                        if mch != machine_in_line.last():
                            document.add_page_break()

                    if line != list_production_line.last():
                        document.add_page_break()

                response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = 'attachment; filename=machine_data.docx'
                document.save(response)
                return response

            elif file_report == 'excel':
                ## Prepare exporting
                response = HttpResponse(content_type='application/ms-excel')
                response['Content-Disposition'] = 'attachment; filename="Machines.xls"'

                wb = xlwt.Workbook(encoding='utf-8')
                ws = wb.add_sheet('Machines Data', cell_overwrite_ok=True)      # this will make a sheet named Machines Data

                # Sheet header, first row
                row_num = 0
                ws.row(0).height_mismatch = True
                ws.row(0).height = 200*2
                style = xlwt.easyxf('pattern: pattern solid, fore_colour gray25;''font: colour black, bold True;''align: vert centre, horiz centre')

                columns = ['Production Line', 'Machine Type', 'Machine Subtype', 'Machine Name', 'Line Code', 'Asset Code', 'Machine Brand', 'Machine Model', 'Machine Serial',
                           'Load Capacity', 'Load Unit', 'Power (Kwatt/Hour)', 'Machine Hour', 'Installed Date', 'Start Date']

                for col_num in range(len(columns)):
                    ws.col(col_num).width = 4500
                    ws.write(row_num, col_num, columns[col_num], style)             # at 0 row

                # Sheet body, remaining rows
                styles = ['pattern: pattern solid, fore_colour ice_blue;' 'align: vert centre, horiz centre',
                          'pattern: pattern solid, fore_colour ivory;''align: vert centre, horiz centre',
                          'pattern: pattern solid, fore_colour light_green;' 'align: vert centre, horiz centre',
                          'pattern: pattern solid, fore_colour lavender;' 'align: vert centre, horiz centre',
                          'pattern: pattern solid, fore_colour light_turquoise;' 'align: vert centre, horiz centre',
                          'pattern: pattern solid, fore_colour tan;' 'align: vert centre, horiz centre',
                          'pattern: pattern solid, fore_colour rose;' 'align: vert centre, horiz centre']

                queryset = Machine.objects.filter(machine_id__in=list_machine).order_by('line__production_line', 'machine_production_line_code')

                rows = queryset.values_list('line__production_line', 'mch_type__mtype_name', 'sub_type__subtype_name', 'machine_name',
                                            'machine_production_line_code', 'machine_asset_code', 'machine_brand', 'machine_model', 'serial_id',
                                            'machine_load_capacity', 'machine_load_capacity_unit', 'machine_power_use_kwatt_per_hour',
                                            'machine_hour', 'machine_installed_datetime', 'machine_start_use_datetime')

                get_lines = rows.values('line__production_line')

                lines = []
                for l in get_lines:
                    lines.append(l.get('line__production_line'))

                for row in rows:
                    row_num += 1
                    for col_num in range(len(row)):
                        if col_num == 0:
                            for l in range(len(lines)):
                                if l == 0:
                                    findex = lines.index(lines[l])
                                    # Get last index of item in list
                                    lindex = len(lines) - lines[::-1].index(lines[l]) - 1
                                    style = xlwt.easyxf(styles[0])
                                    ws.write_merge(findex+1, lindex+1, 0, 0, lines[l], style)

                                elif lines[l] != lines[l-1]:
                                    styles = styles[1:]+styles[0:1]
                                    findex = lines.index(lines[l])
                                    # Get last index of item in list
                                    lindex = len(lines) - lines[::-1].index(lines[l]) - 1
                                    style = xlwt.easyxf(styles[0])
                                    ws.write_merge(findex+1, lindex+1, 0, 0, lines[l], style)
                        else:
                            if isinstance(row[col_num], datetime.date):
                                date_format = xlwt.XFStyle()
                                date_format.num_format_str = 'dd/mm/yyyy'
                                ws.write(row_num, col_num, row[col_num], date_format)
                            else:
                                ws.write(row_num, col_num, row[col_num])

                wb.save(response)

                return response

        return redirect('machine_manage')

    context = {
        'User_login': User_login, 'user_org': user_org,
        'machine': machine, 'mch_subtype_all': mch_subtype_all,
        'production_line': pd_line, 'mch_type_all': mch_type_all, 'role_and_screen': role_and_screen,
        'select_line_export': select_line_export, 'filter_mch_line': filter_mch_line,
        'mch_menu': mch_menu, 'mch_main_menu': mch_main_menu}
    return render(request, 'machine_management/machine_manage.html', context)


def machine_type(request):
    global User_login
    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='machine_type')
    if not role_and_screen.exists():
        return redirect('signin')
    mch_menu = dict_menu_level[Menu.objects.get(pk='machine_manage_menu')]
    mch_main_menu = Menu.objects.get(pk='machine_type')

    mch_types = Machine_type.objects.all()
    if request.method == "POST":
        if 'Addtype' in request.POST:
            add_type_name = request.POST['add_type']
            add_type_code = request.POST['add_type_code']
            added_by = User_login.username
            now = datetime.datetime.now()
            created_date = now.date()
            if not Machine_type.objects.filter(mtype_code=add_type_code, mtype_name=add_type_name).exists():
                add_new_type = Machine_type.objects.create(
                    mtype_code=add_type_code,
                    mtype_name=add_type_name,
                    create_by=added_by,
                    create_date=created_date,
                )
                add_new_type.save()
                messages.success(request, 'เพิ่มข้อมูล Machine Type เรียบร้อยแล้ว')
            else:
                messages.error(request, 'การเพิ่มข้อมูล Machine Type ล้มเหลว กรุณากด Add New Machine Type ใหม่อีกครั้ง')

        elif 'Edittype' in request.POST:
            edit_type = Machine_type.objects.get(mtype_id=request.POST['Edittype'])
            if Machine_type.objects.filter(mtype_code=request.POST['set_type_code']).exists() and edit_type.mtype_code != request.POST['set_type_code']:
                messages.error(request, 'การแก้ไขไม่สำเร็จ เนื่องจาก Machine Code นี้มีอยู่แล้วในระบบ')
            else:
                edit_type.mtype_code = request.POST['set_type_code']
                edit_type.mtype_name = request.POST['set_mch_type']
                edit_type.last_update_by = User_login.username
                now = datetime.datetime.now()
                edit_type.last_update_date = now.date()
                edit_type.save()
                messages.success(request, 'แก้ไขข้อมูล Machine Type สำเร็จ')
        elif 'Delete_type' in request.POST:
            del_type = request.POST['Delete_type']
            typeid = Machine_type.objects.get(mtype_id=del_type)
            typeid.delete()
            messages.success(request, "ลบรายการสำเร็จ")

        return redirect('machine_type')

    context = {'mch_types': mch_types, 'User_login': User_login, 'role_and_screen': role_and_screen,
               'mch_menu': mch_menu, 'mch_main_menu': mch_main_menu}
    return render(request, 'machine_management/machine_type.html', context)


def machine_subtype(request):
    global User_login
    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='machine_sub_type')
    if not role_and_screen.exists():
        return redirect('signin')

    mch_menu = dict_menu_level[Menu.objects.get(pk='machine_manage_menu')]
    mch_main_menu = Menu.objects.get(pk='machine_sub_type')

    mch_subtype = Machine_subtype.objects.all()
    mch_type_all = Machine_type.objects.all()
    if request.method == "POST":
        if 'AddSubtype' in request.POST:
            add_subtype_name = request.POST['add_subtype']
            add_subtype_code = request.POST['add_subtype_code']
            added_by = User_login.username
            now = datetime.datetime.now()
            created_date = now.date()
            add_type = request.POST['add_type']

            if not Machine_subtype.objects.filter(subtype_code=add_subtype_code, subtype_name=add_subtype_name,
                                                  mch_type_id=add_type).exists():
                add_new_subtype = Machine_subtype.objects.create(
                    subtype_code=add_subtype_code,
                    subtype_name=add_subtype_name,
                    create_by=added_by,
                    create_date=created_date,
                    mch_type_id=add_type
                )
                add_new_subtype.save()
                messages.success(request, 'เพิ่มข้อมูล Machine Subtype เรียบร้อยแล้ว')
            else:
                messages.error(request, 'การเพิ่มข้อมูล Machine Subtype ล้มเหลว กรุณากด Add New Machine Subtype ใหม่อีกครั้ง')

        elif 'EditSubtype' in request.POST:
            edit_subtype = Machine_subtype.objects.get(subtype_id=request.POST['EditSubtype'])
            edit_subtype.subtype_name = request.POST['set_subtype']
            edit_subtype.mch_type_id = request.POST['select_type']
            edit_subtype.last_update_by = User_login.username
            now = datetime.datetime.now()
            edit_subtype.last_update_date = now.date()
            if not Machine_subtype.objects.filter(subtype_name=request.POST['set_subtype'],
                                                  mch_type_id=request.POST['select_type']).exists():
                edit_subtype.save()
                messages.success(request, 'แก้ไขข้อมูล Machine Subtype สำเร็จ')
            elif edit_subtype.subtype_name == request.POST['set_subtype'] and edit_subtype.mch_type_id == request.POST['select_type']:
                messages.success(request, 'ไม่มีการ Update รายการใหม่')
            else:
                messages.error(request, 'การแก้ข้อมูล Machine Subtype ไม่ถูกต้อง กรุณาแก้ไขใหม่อีกครั้ง')
        elif 'DeleteSubtype' in request.POST:
            del_subtype = request.POST['DeleteSubtype']
            subtypeid = Machine_subtype.objects.get(subtype_id=del_subtype)
            subtypeid.delete()

        return redirect('machine_subtype')

    context = {'subtypes': mch_subtype, 'mch_type_all': mch_type_all, 'User_login': User_login,
               'role_and_screen': role_and_screen, 'mch_menu': mch_menu, 'mch_main_menu': mch_main_menu}
    return render(request, 'machine_management/machine_subtype.html', context)


def home_machine(request, line):
    global User_login, UserRole, dict_menu_level
    if not Role_Screen.objects.filter(role=UserRole).exists():
        return redirect('signin')
    product_line = Production_line.objects.filter(pid=line)
    products = Product.objects.filter(line__in=product_line)
    machine_line = Machine.objects.filter(line__in=product_line)
    context = {'User_login': User_login, 'UserRole': UserRole, 'dict_menu_level': dict_menu_level.items(),
               'machine_line': machine_line, 'product_line': product_line, 'products': products}
    return render(request, 'home/home_machine.html', context)


def machine_details(request, line, machine):
    global User_login, UserRole, dict_menu_level
    if not Role_Screen.objects.filter(role=UserRole).exists():
        return redirect('signin')
    machine = Machine.objects.filter(machine_id=machine, line_id=line)
    spare_part_of_mch = Machine_sparepart.objects.filter(machine_id__in=machine)
    context = {'User_login': User_login, 'UserRole': UserRole, 'dict_menu_level': dict_menu_level.items(),
               'machine': machine, 'spare_part_of_mch': spare_part_of_mch}
    return render(request, 'home/machine_details.html', context)


def spare_part_manage(request):
    global User_login
    role_and_screen = Role_Screen.objects.filter(role=UserRole, screen_id='spare_part_manage')
    if not role_and_screen.exists():
        return redirect('signin')

    sp_menu = dict_menu_level[Menu.objects.get(pk='spare_part_menu')]
    sp_main_menu = Menu.objects.get(pk='spare_part_manage')

    spare_part_all = Spare_part.objects.all()
    spare_part_group_all = Spare_part_group.objects.all()
    if request.method == 'POST':
        if 'add_spare_part' in request.POST:
            if request.POST['id_sp_type'] == "0" or request.POST['id_sp_subtype'] == "0" or request.POST['id_sp_group'] == "0":
                messages.error(request, "ทำรายการไม่สำเร็จ กรุณาระบุกลุ่มอะไหล่ ประเภทอะไหล่ และชนิดอะไหล่")
            else:
                spare_part = Spare_part.objects.create(spare_part_name=request.POST['add_sp_name'],
                                                       spare_part_model=request.POST['add_sp_model'],
                                                       spare_part_brand=request.POST['add_sp_brand'],
                                                       service_life=request.POST['add_service_life'] if request.POST['add_service_life'] != "" else None,
                                                       service_plan_life=request.POST['add_service_plan_life'] if request.POST['add_service_plan_life'] != "" else None,
                                                       spare_part_detail=request.POST['add_detail'],
                                                       spare_part_group_id=request.POST['id_sp_group'],
                                                       spare_part_type_id=request.POST['id_sp_type'],
                                                       spare_part_sub_type_id=request.POST['id_sp_subtype'],
                                                       create_by=User_login.username,
                                                       create_date=datetime.date.today(),
                                                       spare_part_active=True)
                spare_part.save()
                messages.success(request, "สร้างรายการสำเร็จ")
        elif 'edit_spare_part' in request.POST:
            spare_part = Spare_part.objects.get(pk=request.POST['edit_spare_part'])
            spare_part.spare_part_name = request.POST['set_sp_name']
            spare_part.spare_part_model = request.POST['set_sp_model']
            spare_part.spare_part_brand = request.POST['set_sp_brand']
            spare_part.service_life = request.POST['set_service_life'] if request.POST['set_service_life'] != "" else None
            spare_part.service_plan_life = request.POST['set_service_plan_life'] if request.POST['set_service_plan_life'] != "" else None
            spare_part.spare_part_detail = request.POST['set_detail']
            spare_part.last_update_by = User_login.username
            spare_part.last_update_date = datetime.date.today()
            spare_part.spare_part_active = request.POST.get('set_sp_status', False)
            spare_part.save()
            messages.success(request, "แก้ไขและบันทึกรายการสำเร็จ")
        elif 'delete_spare_part' in request.POST:
            spare_part = Spare_part.objects.get(pk=request.POST['delete_spare_part'])
            spare_part.delete()
            messages.success(request, "ลบรายการสำเร็จ")

        return redirect('spare_part_manage')

    context = {'User_login': User_login, 'spare_part_all': spare_part_all, 'spare_part_group_all': spare_part_group_all,
               'role_and_screen': role_and_screen, 'sp_menu': sp_menu, 'sp_main_menu': sp_main_menu}
    return render(request, 'spare_part_management/spare_part_manage.html', context)


def spare_part_subtype(request):
    global User_login
    role_and_screen = Role_Screen.objects.filter(role=UserRole, screen_id='spare_part_subtype')
    if not role_and_screen.exists():
        return redirect('signin')

    sp_menu = dict_menu_level[Menu.objects.get(pk='spare_part_menu')]
    sp_main_menu = Menu.objects.get(pk='spare_part_subtype')

    spare_part_subtype_all = Spare_part_sub_type.objects.all()
    spare_part_group_all = Spare_part_group.objects.all()
    if request.method == 'POST':
        if 'add_spare_part_subtype' in request.POST:
            if request.POST['select_sp_type'] != "0":
                sp_subtype = Spare_part_sub_type.objects.create(
                    spare_part_sub_type_code=request.POST['add_sp_subtype_code'],
                    spare_part_sub_type_name=request.POST['add_sp_subtype_name'],
                    spare_part_type_id=request.POST['select_sp_type'],
                    create_by=User_login.username,
                    create_date=datetime.date.today())
                sp_subtype.save()
                messages.success(request, "การเพิ่มรายการชนิดอะไหล่เสร็จสมบูรณ์")
            else:
                messages.error(request, "การทำรายการไม่สำเร็จ เนื่องจากกรอกประเภทอะไหล่ไม่ถูกต้อง")
        elif 'edit_spare_part_subtype' in request.POST:
            sp_subtype = Spare_part_sub_type.objects.get(pk=request.POST['edit_spare_part_subtype'])
            sp_subtype.spare_part_sub_type_name = request.POST['set_sp_suptype_name']
            sp_subtype.last_update_by = User_login.username
            sp_subtype.last_update_date = datetime.date.today()
            sp_subtype.save()
            messages.success(request, "การแก้ไขรายการชนิดอะไหล่เสร็จสมบูรณ์")
        elif 'delete_spare_part_subtype' in request.POST:
            sp_subtype = Spare_part_sub_type.objects.get(pk=request.POST['delete_spare_part_subtype'])
            sp_subtype.delete()
            messages.success(request, "การลบรายการชนิดอะไหล่เสร็จสมบูรณ์")

        return redirect('spare_part_subtype')

    context = {'User_login': User_login, 'spare_part_subtype_all': spare_part_subtype_all,
               'spare_part_group_all': spare_part_group_all, 'role_and_screen': role_and_screen,
               'sp_menu': sp_menu, 'sp_main_menu': sp_main_menu}
    return render(request, 'spare_part_management/spare_part_subtype.html', context)


def spare_part_type(request):
    global User_login
    role_and_screen = Role_Screen.objects.filter(role=UserRole, screen_id='spare_part_type')
    if not role_and_screen.exists():
        return redirect('signin')

    sp_menu = dict_menu_level[Menu.objects.get(pk='spare_part_menu')]
    sp_main_menu = Menu.objects.get(pk='spare_part_type')

    spare_part_group_all = Spare_part_group.objects.all()
    sp_type_all = Spare_part_type.objects.all()
    if request.method == "POST":
        if 'add_spare_part_type' in request.POST:
            spare_type = Spare_part_type.objects.create(spare_part_type_code=request.POST['add_sp_type_code'],
                                                        spare_part_type_name=request.POST['add_sp_type_name'],
                                                        create_by=User_login.username,
                                                        create_date=datetime.date.today(),
                                                        spare_part_group_id=request.POST['select_sp_group'])
            spare_type.save()
            messages.success(request, "สร้างรายการสำเร็จ")
        elif 'edit_spare_part_type' in request.POST:
            spare_type = Spare_part_type.objects.get(pk=request.POST['edit_spare_part_type'])
            spare_type.spare_part_type_name = request.POST['set_sp_name']
            spare_type.last_update_by = User_login.username
            spare_type.last_update_date = datetime.date.today()
            spare_type.save()
            messages.success(request, "แก้ไขและบันทึกรายการสำเร็จ")
        elif 'delete_spare_part' in request.POST:
            spare_type = Spare_part_type.objects.get(pk=request.POST['delete_spare_part'])
            spare_type.delete()
            messages.success(request, "ลบรายการสำเร็จ")

        return redirect('spare_part_type')

    context = {'User_login': User_login, 'sp_type_all': sp_type_all, 'spare_part_group_all': spare_part_group_all,
               'role_and_screen': role_and_screen, 'sp_menu': sp_menu, 'sp_main_menu': sp_main_menu}
    return render(request, 'spare_part_management/spare_part_type.html', context)


def spare_part_group(request):
    global User_login
    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='spare_part_group')
    if not role_and_screen.exists():
        return redirect('signin')

    sp_menu = dict_menu_level[Menu.objects.get(pk='spare_part_menu')]
    sp_main_menu = Menu.objects.get(pk='spare_part_group')

    sp_group_all = Spare_part_group.objects.all()
    if request.method == "POST":
        if 'add_spare_part_group' in request.POST:
            spare_group = Spare_part_group.objects.create(spare_part_group_code=request.POST['add_sp_group_code'],
                                                        spare_part_group_name=request.POST['add_sp_group_name'],
                                                        create_by=User_login.username,
                                                        create_date=datetime.date.today())
            spare_group.save()
            messages.success(request, "สร้างรายการสำเร็จ")
        elif 'edit_spare_part_group' in request.POST:
            spare_group = Spare_part_group.objects.get(pk=request.POST['edit_spare_part_group'])
            spare_group.spare_part_group_name = request.POST['set_sp_group_name']
            spare_group.last_update_by = User_login.username
            spare_group.last_update_date = datetime.date.today()
            spare_group.save()
            messages.success(request, "แก้ไขและบันทึกรายการสำเร็จ")
        elif 'delete_spare_part' in request.POST:
            spare_group = Spare_part_group.objects.get(pk=request.POST['delete_spare_part'])
            spare_group.delete()
            messages.success(request, "ลบรายการสำเร็จ")

        return redirect('spare_part_group')

    context = {'User_login': User_login, 'sp_group_all': sp_group_all, 'role_and_screen': role_and_screen,
               'sp_menu': sp_menu, 'sp_main_menu': sp_main_menu}
    return render(request, 'spare_part_management/spare_part_group.html', context)


def machine_and_spare_part(request):
    global User_login
    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='machine_spare_part')
    if not role_and_screen.exists():
        return redirect('signin')

    mch_menu = dict_menu_level[Menu.objects.get(pk='machine_manage_menu')]
    mch_main_menu = Menu.objects.get(pk='machine_and_spare_part')

    dict_mch_sp = {}
    user_org = User_login.org.org_line.all()
    machine = Machine.objects.filter(line__in=user_org)
    mch_and_sp_all = Machine_sparepart.objects.filter(machine__in=machine)
    spare_part_group_all = Spare_part_group.objects.all()
    spare_part_all = Spare_part.objects.all()

    if request.method == "POST":
        if "add_mch_and_sp" in request.POST:
            mch_sp_check_exists = Machine_sparepart.objects.filter(machine_id=request.POST["add_mch_and_sp"], spare_part_id=request.POST["select_sp_name"])

            if request.POST["select_sp_name"] == "0":
                messages.error(request, "การทำรายการผิดพลาด กรุณาระบุอะไหล่")
            elif not mch_sp_check_exists.exists():
                mch_and_sp = Machine_sparepart.objects.create(machine_id=request.POST["add_mch_and_sp"],
                                                              spare_part_id=request.POST["select_sp_name"])
                mch_and_sp.save()
            else:
                messages.error(request, "การทำรายการผิดพลาด เครื่องจักรนี้ มีอะไหล่นี้แล้ว")
            return redirect('machine_and_spare_part')

        elif "delete_spare_part" in request.POST:
            mch_and_sp = Machine_sparepart.objects.filter(machine_id=request.POST['delete_spare_part'],
                                                          spare_part_id=request.POST['select_delete_spare_part'])
            mch_and_sp.delete()
            return redirect('machine_and_spare_part')

        elif "setting" in request.POST:
            mch_and_sp = Machine_sparepart.objects.get(machine_id=request.POST['select_machine'],
                                                       spare_part_id=request.POST['select_spare_part'])

            mch_and_sp.last_mtnchng_hour = request.POST['last_mtn_change'] if request.POST['last_mtn_change'] != "" else None
            mch_and_sp.mtnchng_life_hour = request.POST['life_mtn_hour'] if request.POST['life_mtn_hour'] != "" else None
            mch_and_sp.next_mtnchng_hour = request.POST['next_mtn_change'] if request.POST['next_mtn_change'] != "" else None
            mch_and_sp.last_mtnchk_hour = request.POST['last_mtn_check'] if request.POST['last_mtn_check'] != "" else None
            mch_and_sp.mtnchk_life_hour = request.POST['life_check_hour'] if request.POST['life_check_hour'] != "" else None
            mch_and_sp.next_mtnchk_hour = request.POST['next_mtn_check'] if request.POST['next_mtn_check'] != "" else None
            mch_and_sp.save()
            return redirect('machine_and_spare_part')

        elif 'prd_mch' in request.POST:
            if request.POST['production_line'] == "0":
                machine = Machine.objects.filter(line__in=user_org)
            elif request.POST['production_line'] != "0" and request.POST['machine'] == "":
                machine = Machine.objects.filter(line__in=request.POST['production_line'])
            elif request.POST['production_line'] != "0" and request.POST['machine'] != 0:
                machine = Machine.objects.filter(machine_id=request.POST['machine'])
            mch_and_sp_all = Machine_sparepart.objects.filter(machine__in=machine)

    for mch in machine:
        dict_mch_sp[mch] = []
    for mch_sp in mch_and_sp_all:
        dict_mch_sp[mch_sp.machine].append(mch_sp.spare_part)

    context = {'User_login': User_login,
               'dict_mch_sp': dict_mch_sp, 'spare_part_all': spare_part_all,
               'role_and_screen': role_and_screen, 'spare_part_group_all': spare_part_group_all, 'machine_all': machine,
               'mch_menu': mch_menu, 'mch_main_menu': mch_main_menu, 'line_of_user': user_org}
    return render(request, 'machine_management/machine&spare_part.html', context)


def maintenance_assign(request):

    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='maintenance_assign')
    if not role_and_screen.exists():
        return redirect('signin')

    mtn_menu = dict_menu_level[Menu.objects.get(pk='preventive_data')]
    mtn_main_menu = Menu.objects.get(pk='maintenance_assign')

    maintenance_job_gen = Maintenance_job.objects.all()
    if request.method == "POST":
        if "assign_submit" in request.POST:
            if request.POST.getlist('assign_list[]') != "":
                for job_pk in request.POST.getlist('assign_list[]'):
                    job = Maintenance_job.objects.get(pk=job_pk)
                    if job.job_status in ['รอการอนุมัติงาน', 'งานเสร็จสิ้น', 'ปิดงาน']:
                        messages.error(request, 'ไม่สามารถมอบหมายงานได้ เนื่องจากงานอยู่ในสถานะรอการอนุมัติงานแล้วหรืองานเสร็จสิ้น')
                        continue
                    try:
                        job.job_assign_user = User_login
                        job.job_response_user = User.objects.get(pk=request.POST['user_response']) if request.POST['user_response'] != "" else None
                        job.job_assign_date = datetime.date.today()
                        job.job_status = "รอการดำเนินงาน" if request.POST['user_response'] != "" else "รอการมอบหมาย"
                        job.save()
                    except ObjectDoesNotExist:
                        messages.error(request, "ไม่มีผู้ใช้งานดังกล่าว กรุณากรอกข้อมูลให้ถูกต้อง")
                        return redirect('maintenance_assign')
            return redirect('maintenance_assign')

        elif "set_maintenance_data" in request.POST:
            mch_and_sp = Machine_sparepart.objects.get(pk=request.POST['set_maintenance_data'])
            mch_and_sp.last_mtnchng_hour = request.POST['last_mtn_change'] if request.POST['last_mtn_change'] != "" else None
            mch_and_sp.mtnchng_life_hour = request.POST['life_mtn_hour'] if request.POST['life_mtn_hour'] != "" else None
            mch_and_sp.next_mtnchng_hour = request.POST['next_mtn_change'] if request.POST['next_mtn_change'] != "" else None
            mch_and_sp.last_mtnchk_hour = request.POST['last_mtn_check'] if request.POST['last_mtn_check'] != "" else None
            mch_and_sp.mtnchk_life_hour = request.POST['life_check_hour'] if request.POST['life_check_hour'] != "" else None
            mch_and_sp.next_mtnchk_hour = request.POST['next_mtn_check'] if request.POST['next_mtn_check'] != "" else None
            mch_and_sp.save()
            return redirect('maintenance_assign')

        elif "filter" in request.POST:
            if request.POST['filter_status'] == "0":
                maintenance_job_gen = Maintenance_job.objects.all()
            elif request.POST['filter_status'] == "1":
                maintenance_job_gen = Maintenance_job.objects.filter(job_status="รอการมอบหมาย")
            elif request.POST['filter_status'] == "2":
                maintenance_job_gen = Maintenance_job.objects.exclude(job_status="รอการมอบหมาย")

    context = {'User_login': User_login, 'maintenance_job_gen': maintenance_job_gen,
               'mtn_menu': mtn_menu, 'mtn_main_menu': mtn_main_menu}
    return render(request, 'maintenance/maintenance_assign.html', context)


def machine_capacity(request):
    global User_login, UserRole
    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='machine_capacity')
    if not role_and_screen.exists():
        return redirect('signin')

    mch_menu = dict_menu_level[Menu.objects.get(pk='machine_manage_menu')]
    mch_main_menu = Menu.objects.get(pk='machine_capacity')

    if request.method == "POST":
        if 'Add_machine_capacity' in request.POST:
            mch_capacity = Machine_capacity.objects.filter(machine=request.POST['add_mch'],
                                                           product=request.POST['add_product'])
            if not mch_capacity.exists():
                create_mch_capacity = Machine_capacity.objects.create(
                    machine_id=request.POST['add_mch'],
                    product_id=request.POST['add_product'],
                    rm_name=request.POST['add_rm_name'],
                    rm_batch_size=request.POST['add_rm_batch_size'],
                    rm_unit=request.POST['add_rm_batch_unit'],
                    fg_batch_size=request.POST['add_fg_batch_size'],
                    fg_batch_time=request.POST['add_fg_batch_time'],
                    fg_capacity=request.POST['add_fg_capacity']
                )
                create_mch_capacity.save()
            else:
                messages.info(request, 'รายการเครื่องจักรและผลิตภัณฑ์นี้มีข้อมูลแล้ว ไม่สามารถทำการเพิ่มรายการได้')
        elif 'Edit_mch_capacity' in request.POST:
            edit_mch_cap = Machine_capacity.objects.get(pk=request.POST['Edit_mch_capacity'])
            edit_mch_cap.rm_name = request.POST['set_rm_name']
            edit_mch_cap.rm_batch_size = request.POST['set_rm_batch_size']
            edit_mch_cap.rm_unit = request.POST['set_rm_batch_unit']
            edit_mch_cap.fg_batch_size = request.POST['set_fg_batch_size']
            edit_mch_cap.fg_batch_time = request.POST['set_fg_batch_time']
            edit_mch_cap.fg_capacity = request.POST['set_fg_capacity']
            edit_mch_cap.save()
        elif 'delete_machine_capacity' in request.POST:
            delete_machine_capacity = Machine_capacity.objects.get(pk=request.POST['delete_machine_capacity'])
            delete_machine_capacity.delete()

        return redirect('machine_capacity')

    user_org = User_login.org.org_line.all()
    machine = Machine.objects.filter(line__in=user_org)
    mch_capacity_all = Machine_capacity.objects.filter(machine__in=machine)
    production_line = Production_line.objects.all()
    context = {'User_login': User_login, 'mch_capacity_all': mch_capacity_all, 'production_line': production_line, 'role_and_screen': role_and_screen,
               'user_org': user_org, 'mch_menu': mch_menu, 'mch_main_menu': mch_main_menu}
    return render(request, 'machine_management/machine_capacity.html', context)


def spare_part_and_machine(request):
    global User_login, UserRole
    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='machine_capacity')
    if not role_and_screen.exists():
        return redirect('signin')

    sp_menu = dict_menu_level[Menu.objects.get(pk='spare_part_menu')]
    sp_main_menu = Menu.objects.get(pk='spare_part_and_machine')

    dict_mch_sp = {}
    user_org = User_login.org.org_line.all()
    machine = Machine.objects.filter(line__in=user_org)
    sp_and_mch_all = Machine_sparepart.objects.filter(machine__in=machine)
    spare_part_all = Spare_part.objects.all()
    machine_type_all = Machine_type.objects.all()
    spare_part_group_all = Spare_part_group.objects.all()

    if request.method == "POST":
        if "add_sp_and_mch" in request.POST:
            try:
                sp_and_mch = Machine_sparepart.objects.create(machine_id=request.POST['select_mch'], spare_part_id=request.POST['add_sp_and_mch'])
                sp_and_mch.save()
                messages.success(request, 'บันทึกรายการเครื่องจักรออกจากอะไหล่สำเร็จ')
            except ObjectDoesNotExist:
                messages.error(request, 'ลบรายการเครื่องจักรออกจากอะไหล่ไม่สำเร็จ')
            return redirect('spare_part_and_machine')
        elif "delete_machine" in request.POST:
            try:
                sp_and_mch = Machine_sparepart.objects.filter(machine_id=request.POST['select_delete_machine'], spare_part_id=request.POST['delete_machine'])
                sp_and_mch.delete()
                messages.success(request, 'ลบรายการเครื่องจักรออกจากอะไหล่สำเร็จ')
            except ObjectDoesNotExist:
                messages.error(request, 'ลบรายการเครื่องจักรออกจากอะไหล่ไม่สำเร็จ')
            return redirect('spare_part_and_machine')
        elif "filter_spare_part" in request.POST:
            if request.POST['sp_group'] == "0":
                pass
            elif request.POST['sp_group'] != "0" and request.POST['sp_type'] == "0":
                spare_part_all = Spare_part.objects.filter(spare_part_group_id=request.POST['sp_group'])
            elif request.POST['sp_group'] != "0" and request.POST['sp_type'] != 0 and request.POST['sp_subtype'] == "":
                spare_part_all = Spare_part.objects.filter(spare_part_type_id=request.POST['sp_type'])
            elif request.POST['sp_group'] != "0" and request.POST['sp_type'] != 0 and request.POST['sp_subtype'] != 0:
                spare_part_all = Spare_part.objects.filter(spare_part_sub_type_id=request.POST['sp_subtype'])
            sp_and_mch_all = Machine_sparepart.objects.filter(machine__in=machine, spare_part__in=spare_part_all)

    for sp in spare_part_all:
        dict_mch_sp[sp] = []
    for sp_mch in sp_and_mch_all:
        dict_mch_sp[sp_mch.spare_part].append(sp_mch.machine)

    context = {'role_and_screen': role_and_screen, 'spare_part_all': spare_part_all, 'dict_mch_sp': dict_mch_sp, 'user_org': user_org,
               'machine_type_all': machine_type_all, 'User_login': User_login, 'sp_menu': sp_menu, 'sp_main_menu': sp_main_menu,
               'spare_part_group_all': spare_part_group_all}
    return render(request, 'spare_part_management/spare_and_machine.html', context)


def maintenance_data(request):
    global User_login
    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='maintenance_data')
    if not role_and_screen.exists():
        return redirect('signin')

    mtn_menu = dict_menu_level[Menu.objects.get(pk='preventive_data')]
    mtn_main_menu = Menu.objects.get(pk='maintenance_data')

    line_of_user = User_login.org.org_line.all()
    mch_sp_all = Machine_sparepart.objects.filter(machine__line__in=line_of_user)
    if request.method == 'POST':

        if 'prd_mch' in request.POST:
            if request.POST['production_line'] == "0":
                mch_sp_all = Machine_sparepart.objects.filter(machine__line__in=line_of_user)
            elif request.POST['production_line'] != "0" and request.POST['machine'] == "":
                mch_sp_all = Machine_sparepart.objects.filter(machine__line__in=request.POST['production_line'])
            elif request.POST['production_line'] != "0" and request.POST['machine'] != 0:
                mch_sp_all = Machine_sparepart.objects.filter(machine_id=request.POST['machine'])

        elif "set_maintenance_data" in request.POST:
            mch_and_sp = Machine_sparepart.objects.get(pk=request.POST['set_maintenance_data'])
            mch_and_sp.last_mtnchng_hour = request.POST['last_mtn_change'] if request.POST['last_mtn_change'] != "" else None
            mch_and_sp.mtnchng_life_hour = request.POST['life_mtn_hour'] if request.POST['life_mtn_hour'] != "" else None
            mch_and_sp.next_mtnchng_hour = request.POST['next_mtn_change'] if request.POST['next_mtn_change'] != "" else None
            mch_and_sp.last_mtnchk_hour = request.POST['last_mtn_check'] if request.POST['last_mtn_check'] != "" else None
            mch_and_sp.mtnchk_life_hour = request.POST['life_check_hour'] if request.POST['life_check_hour'] != "" else None
            mch_and_sp.next_mtnchk_hour = request.POST['next_mtn_check'] if request.POST['next_mtn_check'] != "" else None
            mch_and_sp.save()
            return redirect('maintenance_data')

    context = {'User_login': User_login, 'line_of_user': line_of_user, 'mch_sp_all': mch_sp_all,
               'menu_job': dict_menu_level[Menu.objects.get(menu_id='preventive_data')],
               'mtn_menu': mtn_menu, 'mtn_main_menu': mtn_main_menu}
    return render(request, 'maintenance/maintenance_data.html', context)


def maintenance_report(request):
    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='maintenance_report')
    if not role_and_screen.exists():
        return redirect('signin')

    mtn_menu = dict_menu_level[Menu.objects.get(pk='preventive_data')]
    mtn_main_menu = Menu.objects.get(pk='maintenance_report')

    job = Maintenance_job.objects.filter(Q(job_response_user_id=User_login.username) | Q(job_assign_user_id=User_login.username))
    if request.method == "POST":
        if 'report_submit' in request.POST:
            mtn_report = Maintenance_job.objects.get(pk=request.POST['report_submit'])
            mch_sp = Machine_sparepart.objects.get(pk=mtn_report.job_mch_sp_id)
            mtn_report.job_mtn_type = request.POST['mtn_type']
            mtn_report.job_result_type = request.POST['mtn_result']
            mtn_report.job_result_description = request.POST['mtn_result_description'] if request.POST['mtn_result_description'] != "" else None
            mtn_report.job_fix_plan_hour = request.POST['chang_life_hour_sp'] if request.POST['chang_life_hour_sp'] != "" else request.POST['chang_life_hour_pv']
            mtn_report.job_mch_hour = request.POST['machine_hour_sp'] if request.POST['machine_hour_sp'] != "" else request.POST['machine_hour_pv']
            mtn_report.job_plan_hour = request.POST['check_life_hour_sp'] if request.POST['check_life_hour_sp'] != "" else request.POST['check_life_hour_pv']
            mtn_report.job_report_date = datetime.datetime.today()
            mtn_report.problem_cause = request.POST['problem_cause']
            mtn_report.corrective_action = request.POST['corrective_action']
            mtn_report.after_repair = request.POST['after_repair']
            mtn_report.job_status = "รอการอนุมัติงาน"
            mtn_report.is_approve = False
            mtn_report.is_report = True
            mtn_report.equipment_code1 = request.POST['equipment_code1'] if request.POST['equipment_code1'] != "" else None
            mtn_report.equipment_code2 = request.POST['equipment_code2'] if request.POST['equipment_code2'] != "" else None
            mtn_report.equipment_code3 = request.POST['equipment_code3'] if request.POST['equipment_code3'] != "" else None
            mtn_report.equipment_detail1 = request.POST['equipment_detail1'] if request.POST['equipment_detail1'] != "" else None
            mtn_report.equipment_detail2 = request.POST['equipment_detail2'] if request.POST['equipment_detail2'] != "" else None
            mtn_report.equipment_detail3 = request.POST['equipment_detail3'] if request.POST['equipment_detail3'] != "" else None
            mtn_report.equipment_quantity1 = request.POST['equipment_quantity1'] if request.POST['equipment_quantity1'] != "" else None
            mtn_report.equipment_quantity2 = request.POST['equipment_quantity2'] if request.POST['equipment_quantity2'] != "" else None
            mtn_report.equipment_quantity3 = request.POST['equipment_quantity3'] if request.POST['equipment_quantity3'] != "" else None
            mtn_report.equipment_note1 = request.POST['equipment_note1'] if request.POST['equipment_note1'] != "" else None
            mtn_report.equipment_note2 = request.POST['equipment_note2'] if request.POST['equipment_note2'] != "" else None
            mtn_report.equipment_note3 = request.POST['equipment_note3'] if request.POST['equipment_note3'] != "" else None
            mtn_report.job_remark = request.POST['job_remark'] if request.POST['job_remark'] != "" else None
            mtn_report.estimate_cost = request.POST['estimate_cost'] if request.POST['estimate_cost'] != "" else None
            try:
                mtn_report.save()
                mch_sp.save()
            except ValueError:
                messages.error(request, 'กรุณากรอกชั่วโมงการเปลี่ยนและชั่วโมงการตรวจสอบของอะไหล่')
            return redirect('maintenance_report')

        elif "approve_job" in request.POST:
            if request.POST.get('is_approve', False):
                mtn_report = Maintenance_job.objects.get(pk=request.POST['approve_job'])
                mch_sp = Machine_sparepart.objects.get(pk=mtn_report.job_mch_sp_id)
                mtn_report.is_approve = True
                mtn_report.job_approve_date = datetime.datetime.today()
                mtn_report.job_status = "งานเสร็จสิ้น" if request.POST.get('is_approve', False) else "รอการอนุมัติงาน"
                mch_sp.gen_mtnchng_date = None
                mch_sp.gen_mtnchk_date = None
                machine_model = Machine.objects.get(pk=mch_sp.machine.pk)
                machine_model.machine_hour_last_update = mch_sp.machine.machine_hour
                machine_model.machine_hour = mtn_report.job_mch_hour

                if mtn_report.job_mtn_type in ["change", "repair"]:
                    mch_sp.last_mtnchk_hour = mtn_report.job_mch_hour
                    mch_sp.last_mtnchng_hour = mtn_report.job_mch_hour
                    mch_sp.last_mtnchng_job_id = mtn_report.id
                elif mtn_report.job_mtn_type == "checking":
                    mch_sp.last_mtnchk_hour = mtn_report.job_mch_hour
                    mch_sp.last_mtnchk_job_id = mtn_report.id
                mch_sp.mtnchng_life_hour = mtn_report.job_fix_plan_hour
                mch_sp.mtnchk_life_hour = mtn_report.job_plan_hour
                if mch_sp.mtnchk_life_hour:
                    mch_sp.next_mtnchk_hour = int(mch_sp.last_mtnchk_hour) + int(mch_sp.mtnchk_life_hour)
                if mch_sp.mtnchng_life_hour:
                    mch_sp.next_mtnchng_hour = int(mch_sp.last_mtnchng_hour) + int(mch_sp.mtnchng_life_hour)

                machine_model.save()
                mch_sp.save()
                mtn_report.save()

            else:
                mtn_report = Maintenance_job.objects.get(pk=request.POST['approve_job'])
                mtn_report.job_status = "ไม่ผ่านการอนุมัติ"
                mtn_report.is_report = False
                mtn_report.save()
            return redirect('maintenance_report')

        elif "filter" in request.POST:
            if request.POST['filter_date'] == "0":
                if request.POST['filter_status'] == "0":
                    pass
                elif request.POST['filter_status'] == "1":
                    job = Maintenance_job.objects.filter(Q(job_response_user_id=User_login.username) | Q(job_assign_user_id=User_login.username)).exclude(Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน"))
                elif request.POST['filter_status'] == "2":
                    job = Maintenance_job.objects.filter(Q(job_response_user_id=User_login.username) | Q(job_assign_user_id=User_login.username) | Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน"))
                elif request.POST['filter_status'] == "3":
                    job = Maintenance_job.objects.all()
                elif request.POST['filter_status'] == "4":
                    job = Maintenance_job.objects.exclude(Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน"))
                elif request.POST['filter_status'] == "5":
                    job = Maintenance_job.objects.filter(Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน"))
            elif request.POST['filter_date'] == "1":
                date_today = datetime.date.today()
                if request.POST['filter_status'] == "0":
                    job = Maintenance_job.objects.filter(Q(job_response_user_id=User_login.username) | Q(job_assign_user_id=User_login.username) | Q(job_gen_date=date_today))
                elif request.POST['filter_status'] == "1":
                    job = Maintenance_job.objects.filter(Q(job_response_user_id=User_login.username) | Q(job_assign_user_id=User_login.username) | Q(job_gen_date=date_today)).exclude(Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน"))
                elif request.POST['filter_status'] == "2":
                    job = Maintenance_job.objects.filter(Q(job_response_user_id=User_login.username) | Q(job_assign_user_id=User_login.username) | Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน") | Q(job_gen_date=date_today))
                elif request.POST['filter_status'] == "3":
                    job = Maintenance_job.objects.filter(job_gen_date=date_today)
                elif request.POST['filter_status'] == "4":
                    job = Maintenance_job.objects.exclude(Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน") & Q(job_gen_date=date_today))
                elif request.POST['filter_status'] == "5":
                    job = Maintenance_job.objects.filter(Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน") & Q(job_gen_date=date_today))
            elif request.POST['filter_date'] == "2":
                end_date = datetime.date.today()
                start_date = datetime.date.today()-datetime.timedelta(days=7)
                if request.POST['filter_status'] == "0":
                    job = Maintenance_job.objects.filter(Q(job_response_user_id=User_login.username) | Q(job_assign_user_id=User_login.username) | Q(job_gen_date__range=[start_date, end_date]))
                elif request.POST['filter_status'] == "1":
                    job = Maintenance_job.objects.filter(Q(job_response_user_id=User_login.username) | Q(job_assign_user_id=User_login.username) | Q(job_gen_date__range=[start_date, end_date])).exclude(Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน"))
                elif request.POST['filter_status'] == "2":
                    job = Maintenance_job.objects.filter(Q(job_response_user_id=User_login.username) | Q(job_assign_user_id=User_login.username) | Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน") | Q(job_gen_date__range=[start_date, end_date]))
                elif request.POST['filter_status'] == "3":
                    job = Maintenance_job.objects.filter(job_gen_date__range=[start_date, end_date])
                elif request.POST['filter_status'] == "4":
                    job = Maintenance_job.objects.exclude(Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน") & Q(job_gen_date__range=[start_date, end_date]))
                elif request.POST['filter_status'] == "5":
                    job = Maintenance_job.objects.filter(Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน") & Q(job_gen_date__range=[start_date, end_date]))
            elif request.POST['filter_date'] == "3":
                this_month = datetime.date.today().month
                if request.POST['filter_status'] == "0":
                    job = Maintenance_job.objects.filter(Q(job_response_user_id=User_login.username) | Q(job_assign_user_id=User_login.username) | Q(job_gen_date__month=this_month))
                elif request.POST['filter_status'] == "1":
                    job = Maintenance_job.objects.filter(Q(job_response_user_id=User_login.username) | Q(job_assign_user_id=User_login.username) | Q(job_gen_date__month=this_month)).exclude(Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน"))
                elif request.POST['filter_status'] == "2":
                    job = Maintenance_job.objects.filter(Q(job_response_user_id=User_login.username) | Q(job_assign_user_id=User_login.username) | Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน") | Q(job_gen_date__month=this_month))
                elif request.POST['filter_status'] == "3":
                    job = Maintenance_job.objects.filter(job_gen_date__month=this_month)
                elif request.POST['filter_status'] == "4":
                    job = Maintenance_job.objects.exclude(Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน") & Q(job_gen_date__month=this_month))
                elif request.POST['filter_status'] == "5":
                    job = Maintenance_job.objects.filter(Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน") & Q(job_gen_date__month=this_month))
            elif request.POST['filter_date'] == "4":
                this_year = datetime.date.today().year
                if request.POST['filter_status'] == "0":
                    job = Maintenance_job.objects.filter(Q(job_response_user_id=User_login.username) | Q(job_assign_user_id=User_login.username) | Q(job_gen_date__year=this_year))
                elif request.POST['filter_status'] == "1":
                    job = Maintenance_job.objects.filter(Q(job_response_user_id=User_login.username) | Q(job_assign_user_id=User_login.username) | Q(job_gen_date__year=this_year)).exclude(Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน"))
                elif request.POST['filter_status'] == "2":
                    job = Maintenance_job.objects.filter(Q(job_response_user_id=User_login.username) | Q(job_assign_user_id=User_login.username) | Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน") | Q(job_gen_date__year=this_year))
                elif request.POST['filter_status'] == "3":
                    job = Maintenance_job.objects.filter(job_gen_date__year=this_year)
                elif request.POST['filter_status'] == "4":
                    job = Maintenance_job.objects.exclude(Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน") & Q(job_gen_date__year=this_year))
                elif request.POST['filter_status'] == "5":
                    job = Maintenance_job.objects.filter(Q(job_status="งานเสร็จสิ้น") | Q(job_status="ปิดงาน") & Q(job_gen_date__year=this_year))

        elif 'export' in request.POST:
            if request.POST['people_file'] == "1":
                if request.POST['date_file'] == "1":
                    date_today = datetime.date.today()
                    if request.POST['status_file'] == "1":
                        mtn_job = Maintenance_job.objects.filter(job_response_user=User_login, job_gen_date=date_today).exclude(job_status="งานเสร็จสิ้น")
                    elif request.POST['status_file'] == "2":
                        mtn_job = Maintenance_job.objects.filter(job_response_user=User_login, job_gen_date=date_today, job_status="งานเสร็จสิ้น")
                    elif request.POST['status_file'] == "3":
                        mtn_job = Maintenance_job.objects.filter(job_response_user=User_login, job_gen_date=date_today)
                elif request.POST['date_file'] == "2":
                    end_date = datetime.date.today()
                    start_date = datetime.date.today()-datetime.timedelta(days=7)
                    if request.POST['status_file'] == "1":
                        mtn_job = Maintenance_job.objects.filter(job_response_user=User_login, job_gen_date__range=[start_date, end_date]).exclude(job_status="งานเสร็จสิ้น")
                    elif request.POST['status_file'] == "2":
                        mtn_job = Maintenance_job.objects.filter(job_response_user=User_login, job_gen_date__range=[start_date, end_date], job_status="งานเสร็จสิ้น")
                    elif request.POST['status_file'] == "3":
                        mtn_job = Maintenance_job.objects.filter(job_response_user=User_login, job_gen_date__range=[start_date, end_date])
                elif request.POST['date_file'] == "3":
                    this_month = datetime.date.today().month
                    if request.POST['status_file'] == "1":
                        mtn_job = Maintenance_job.objects.filter(job_response_user=User_login, job_gen_date__month=this_month).exclude(job_status="งานเสร็จสิ้น")
                    elif request.POST['status_file'] == "2":
                        mtn_job = Maintenance_job.objects.filter(job_response_user=User_login, job_gen_date__month=this_month, job_status="งานเสร็จสิ้น")
                    elif request.POST['status_file'] == "3":
                        mtn_job = Maintenance_job.objects.filter(job_response_user=User_login, job_gen_date__month=this_month)
                elif request.POST['date_file'] == "4":
                    this_year = datetime.date.today().year
                    if request.POST['status_file'] == "1":
                        mtn_job = Maintenance_job.objects.filter(job_response_user=User_login, job_gen_date__year=this_year).exclude(job_status="งานเสร็จสิ้น")
                    elif request.POST['status_file'] == "2":
                        mtn_job = Maintenance_job.objects.filter(job_response_user=User_login, job_gen_date__year=this_year, job_status="งานเสร็จสิ้น")
                    elif request.POST['status_file'] == "3":
                        mtn_job = Maintenance_job.objects.filter(job_response_user=User_login, job_gen_date__year=this_year)
            elif request.POST['people_file'] == "2":
                if request.POST['date_file'] == "1":
                    date_today = datetime.date.today()
                    if request.POST['status_file'] == "1":
                        mtn_job = Maintenance_job.objects.filter(job_gen_date=date_today).exclude(job_status="งานเสร็จสิ้น")
                    elif request.POST['status_file'] == "2":
                        mtn_job = Maintenance_job.objects.filter(job_gen_date=date_today, job_status="งานเสร็จสิ้น")
                    elif request.POST['status_file'] == "3":
                        mtn_job = Maintenance_job.objects.filter(job_gen_date=date_today)
                elif request.POST['date_file'] == "2":
                    end_date = datetime.date.today()
                    start_date = datetime.date.today()-datetime.timedelta(days=7)
                    if request.POST['status_file'] == "1":
                        mtn_job = Maintenance_job.objects.filter(job_gen_date__range=[start_date, end_date]).exclude(job_status="งานเสร็จสิ้น")
                    elif request.POST['status_file'] == "2":
                        mtn_job = Maintenance_job.objects.filter(job_gen_date__range=[start_date, end_date], job_status="งานเสร็จสิ้น")
                    elif request.POST['status_file'] == "3":
                        mtn_job = Maintenance_job.objects.filter(job_gen_date__range=[start_date, end_date])
                elif request.POST['date_file'] == "3":
                    this_month = datetime.date.today().month
                    if request.POST['status_file'] == "1":
                        mtn_job = Maintenance_job.objects.filter(job_gen_date__month=this_month).exclude(job_status="งานเสร็จสิ้น")
                    elif request.POST['status_file'] == "2":
                        mtn_job = Maintenance_job.objects.filter(job_gen_date__month=this_month, job_status="งานเสร็จสิ้น")
                    elif request.POST['status_file'] == "3":
                        mtn_job = Maintenance_job.objects.filter(job_gen_date__month=this_month)
                elif request.POST['date_file'] == "4":
                    this_year = datetime.date.today().year
                    if request.POST['status_file'] == "1":
                        mtn_job = Maintenance_job.objects.filter(job_gen_date__year=this_year).exclude(job_status="งานเสร็จสิ้น")
                    elif request.POST['status_file'] == "2":
                        mtn_job = Maintenance_job.objects.filter(job_gen_date__year=this_year, job_status="งานเสร็จสิ้น")
                    elif request.POST['status_file'] == "3":
                        mtn_job = Maintenance_job.objects.filter(job_gen_date__year=this_year)

            document = Document()
            if request.POST['status_file'] == "1":
                document.add_heading(f'งานระบบซ่อมบำรุงเครื่องจักร (ค้างซ่อม)', 0)
            elif request.POST['status_file'] == "2":
                document.add_heading(f'งานระบบซ่อมบำรุงเครื่องจักร (เสร็จสิ้น)', 0)
            elif request.POST['status_file'] == "3":
                document.add_heading(f'งานระบบซ่อมบำรุงเครื่องจักร (ทั้งหมด)', 0)

            for index, job_export in enumerate(mtn_job):
                document.add_heading(f'{index+1}. หมายเลขงานซ่อมบำรุงเครื่องจักร : {job_export.job_no}', level=1)
                if job_export.job_mch_sp.machine.machine_image1:
                    picture_mch = document.add_picture(job_export.job_mch_sp.machine.machine_image1, width=Inches(2))
                    last_paragraph = document.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                document.add_paragraph('ข้อมูลงานซ่อมบำรุงเครื่องจักร', style='List Bullet')
                document.add_paragraph(f'\tวันที่สร้างงานซ่อมบำรุง : {job_export.job_gen_date}')
                if job_export.job_assign_user:
                    document.add_paragraph(f'\tผู้มอบหมายงาน: {job_export.job_assign_user.firstname}  {job_export.job_assign_user.lastname} \tวันที่มอบหมายงาน : {job_export.job_assign_date}')
                if job_export.job_response_user:
                    document.add_paragraph(f'\tผู้รับผิดชอบงาน : {job_export.job_response_user.firstname}  {job_export.job_response_user.lastname} \tวันที่รายงาน : {job_export.job_report_date}')
                p = document.add_paragraph(f'\tสถานะใบแจ้งซ่อมเครื่องจักร : ')
                p.add_run(job_export.job_status).bold = True
                if job_export.job_approve_date:
                    p.add_run(f'\tวันที่อนุมัติงาน : {job_export.job_approve_date}')

                document.add_paragraph('ข้อมูลเครื่องจักรและอะไหล่ที่ซ่อมบำรุง', style='List Bullet')
                document.add_paragraph(f'\tสายการผลิตที่ : {job_export.job_mch_sp.machine.line}')
                document.add_paragraph(f'\tเครื่องจักร : {job_export.job_mch_sp.machine.machine_production_line_code} | {job_export.job_mch_sp.machine.machine_name}')
                document.add_paragraph(f'\tเครื่องจักร : {job_export.job_mch_sp.spare_part.spare_part_group} | {job_export.job_mch_sp.spare_part}')

                document.add_paragraph('ข้อมูลงานและปัญหา', style='List Bullet')
                document.add_paragraph(f'\tสาเหตุของปัญหา : {job_export.problem_cause}')
                document.add_paragraph(f'\tวิธีการแก้ไข : {job_export.corrective_action}')
                document.add_paragraph(f'\tวิธีการดูแลหลังซ่อมบำรุง : {job_export.after_repair}')

                document.add_paragraph('รายการอุปกรณ์ / ชิ้นส่วนในการใช้ซ่อมบำรุง', style='List Bullet')
                if not job_export.equipment_code1 and not job_export.equipment_code2 and not job_export.equipment_code3:
                    document.add_paragraph(f'\tไม่มีอุปกรณ์ที่ใช้')
                    document.add_paragraph(f'\tค่าใช้จ่ายโดยประมาณ : ไม่มี')
                else:
                    records = []
                    if job_export.equipment_code1:
                        columns = [job_export.equipment_code1, job_export.equipment_detail1, job_export.equipment_quantity1, job_export.equipment_note1]
                        records.append(columns)
                    if job_export.equipment_code2:
                        columns = [job_export.equipment_code2, job_export.equipment_detail2, job_export.equipment_quantity2, job_export.equipment_note2]
                        records.append(columns)
                    if job_export.equipment_code3:
                        columns = [job_export.equipment_code3, job_export.equipment_detail3, job_export.equipment_quantity3, job_export.equipment_note3]
                        records.append(columns)
                    table = document.add_table(rows=1, cols=4)
                    table.style = 'Light List Accent 1'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Equipment Code'
                    hdr_cells[1].text = 'Equipment Name'
                    hdr_cells[2].text = 'Quantity'
                    hdr_cells[3].text = 'Note'
                    for code, name, quantity, note in records:
                        row_cells = table.add_row().cells
                        row_cells[0].text = code
                        row_cells[1].text = name
                        row_cells[2].text = str(quantity)
                        row_cells[3].text = note
                    document.add_paragraph(f'')

                if job_export.job_result_type:
                    document.add_paragraph('ข้อมูลการซ่อมบำรุงรักษา', style='List Bullet')
                    document.add_paragraph(f'\tประเภทการซ่อมบำรุงรักษา : {job_export.job_mtn_type}')
                    document.add_paragraph(f'\tผลการซ่อมบำรุง/ตรวจสอบ : {job_export.job_result_type}')
                    document.add_paragraph(f'\tรายละเอียดเพิ่มเติม : {job_export.job_result_description}')
                    document.add_paragraph(f'\tชั่วโมงเครื่องจักรที่บันทึกผล : {job_export.job_mch_hour}')
                    document.add_paragraph(f'\tอายุการใช้การเปลี่ยนของอะไหล่ที่บันทึกผล : {job_export.job_fix_plan_hour}')
                    document.add_paragraph(f'\tอายุการใช้การตรวจสอบของอะไหล่ที่บันทึกผล : {job_export.job_plan_hour}')


                if index+1 != len(mtn_job):
                    document.add_page_break()

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=maintenance_job.docx'
            document.save(response)
            return response

    context = {'User_login': User_login, 'job': job, 'menu_job': dict_menu_level[Menu.objects.get(menu_id='preventive_data')],
               'mtn_menu': mtn_menu, 'mtn_main_menu': mtn_main_menu}
    return render(request, 'maintenance/maintenance_report.html', context)


def machine_hour_update(request):
    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='machine_hour_update')
    if not role_and_screen.exists():
        return redirect('/')

    mtn_menu = dict_menu_level[Menu.objects.get(pk='preventive_data')]
    mtn_main_menu = Menu.objects.get(pk='machine_hour_update')

    if request.method == 'POST':
        if request.POST.get('mch_update[]', False):
            hour_update = set(request.POST.getlist('hour_update'))
            if '' in hour_update:
                hour_update.remove('')
            hour_update = list(hour_update)[0] if len(hour_update) == 1 else 0
            mch_pk = request.POST.getlist('mch_update[]')
            for mch in Machine.objects.filter(pk__in=mch_pk):
                if mch.machine_hour_update_date == datetime.date.today():
                    mch.machine_hour = mch.machine_hour_last_update + int(hour_update)
                else:
                    mch.machine_hour_update_date = datetime.date.today()
                    mch.machine_hour_last_update = mch.machine_hour
                    mch.machine_hour += int(hour_update)
                mch.save()
            return redirect('machine_hour_update')
        else:
            return redirect('machine_hour_update')

    user_org = User_login.org.org_line.all()
    machine_all = Machine.objects.filter(line__in=user_org)

    context = {'User_login': User_login, 'machine_all': machine_all,
               'mtn_menu': mtn_menu, 'mtn_main_menu': mtn_main_menu}
    return render(request, 'maintenance/machine_hour_update.html', context)


def repair_notice(request):
    global User_login
    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='repair_notice')
    if not role_and_screen.exists():
        return redirect('signin')

    repair_menu = dict_menu_level[Menu.objects.get(pk='repair_menu')]
    repair_main_menu = Menu.objects.get(pk='repair_notice')

    line_of_user = User_login.org.org_line.all()
    list_repair_notice = Repair_notice.objects.filter(repairer_user=User_login)
    list_inspect_user = User_and_department.objects.filter(department__department_code=UserLoginDepartment.department_code, is_inspect=True)
    list_approve_user = User_and_department.objects.filter(department__department_code=UserLoginDepartment.department_code, is_approve=True)
    list_department = Department.objects.all()

    if request.method == "POST":
        if "create_repair" in request.POST:
            list_datetime = request.POST['notification_date'].split('-')
            date_no = list_datetime[0][-2:] + list_datetime[1] + list_datetime[2]
            dep_code = str(Department.objects.get(pk=request.POST['department_notifying']).department_code)
            last_no = Repair_notice.objects.filter(repair_no__startswith="RP"+date_no+dep_code).last()
            if not last_no:
                new_rep_number = "RP" + date_no + dep_code + "001"
            else:
                rep_no = last_no.repair_no
                rep_int = int(rep_no.split(date_no+dep_code)[-1])
                new_rep_int = rep_int + 1
                new_rep_number = "RP" + date_no + dep_code + '{:03}'.format(new_rep_int)
            repair_notice_model = Repair_notice.objects.create(
                repair_no=new_rep_number,
                department_notifying_id=request.POST['department_notifying'],
                department_receive_id=request.POST['department_receive'],
                use_date=request.POST['use_date'],
                repairer_user=User.objects.get(pk=request.POST['repairer_user']),
                inspect_user=User.objects.get(pk=request.POST['inspect_user']),
                approve_user=User.objects.get(pk=request.POST['approve_user']),
                machine_id=request.POST['machine'],
                problem_report=request.POST['problem_report'],
                effect_problem=request.POST['effect_problem'],
                notification_date=request.POST['notification_date']
            )
            if repair_notice_model.repairer_user == repair_notice_model.inspect_user and repair_notice_model.repairer_user == repair_notice_model.approve_user:
                repair_notice_model.repair_status = 'รอการรับใบแจ้ง'
                repair_notice_model.is_inspect = True
                repair_notice_model.is_approve = True
            elif repair_notice_model.repairer_user == repair_notice_model.inspect_user:
                repair_notice_model.repair_status = 'รอการอนุมัติ'
                repair_notice_model.is_inspect = True
            else:
                repair_notice_model.repair_status = 'รอการตรวจสอบ'
            repair_notice_model.save()
            messages.success(request, "สร้างรายการสำเร็จ")
            return redirect('repair_notice')

        elif "set_repair" in request.POST:
            repair_notice_model = Repair_notice.objects.get(pk=request.POST['set_repair'])
            repair_notice_model.notification_date = request.POST['set_notification_date']
            repair_notice_model.machine_id = request.POST['set_machine']
            repair_notice_model.department_receive_id = request.POST['set_department_receive']
            repair_notice_model.use_date = request.POST['set_use_date']
            repair_notice_model.problem_report = request.POST['problem_report']
            repair_notice_model.effect_problem = request.POST['effect_problem']
            repair_notice_model.inspect_user = User.objects.get(pk=request.POST['set_inspect_user'])
            repair_notice_model.approve_user = User.objects.get(pk=request.POST['set_approve_user'])
            if repair_notice_model.repairer_user == repair_notice_model.inspect_user and repair_notice_model.repairer_user == repair_notice_model.approve_user:
                repair_notice_model.repair_status = 'รอการรับใบแจ้ง'
                repair_notice_model.is_inspect = True
                repair_notice_model.is_approve = True
                repair_notice_model.inspect_remark = None
                repair_notice_model.approve_remark = None
            elif repair_notice_model.repairer_user == repair_notice_model.inspect_user:
                repair_notice_model.repair_status = 'รอการอนุมัติ'
                repair_notice_model.is_inspect = True
                repair_notice_model.inspect_remark = None
            else:
                repair_notice_model.repair_status = 'รอการตรวจสอบ'
                repair_notice_model.is_inspect = None
                repair_notice_model.is_approve = None
                repair_notice_model.is_receive = None
                repair_notice_model.inspect_remark = None
                repair_notice_model.approve_remark = None
                repair_notice_model.receive_remark = None

            repair_notice_model.save()
            return redirect('repair_notice')

        elif "delete_repair_notice" in request.POST:
            repair_notice_model = Repair_notice.objects.get(pk=request.POST['delete_repair_notice'])
            repair_notice_model.delete()
            return redirect('repair_notice')

        elif "repair_close" in request.POST:
            repair_notice_model = Repair_notice.objects.get(pk=request.POST['repair_close'])
            if 'is_close' == request.POST['is_close']:
                job_status_list = repair_notice_model.maintenance_jobs.values_list('job_status', flat=True)
                if not all(job in ['ปิดงาน', 'งานเสร็จสิ้น'] for job in job_status_list):
                    messages.error(request, 'มีบางงานของใบแจ้งซ่อมกำลังดำเนินการอยู่')
                    return redirect('repair_notice')
                repair_notice_model.repair_status = 'ปิดใบแจ้งซ่อม'
                repair_notice_model.is_close = True
                repair_notice_model.repair_close_date = datetime.datetime.today()

            elif 'is_cancel' == request.POST['is_close']:
                repair_notice_model.repair_status = 'ยกเลิกใบแจ้ง'
                repair_notice_model.is_close = False
                repair_notice_model.is_cancel = True

            messages.success(request, 'บันทึกข้อมูลสำเร็จ')
            repair_notice_model.close_remark = request.POST['close_remark'] if request.POST['close_remark'] != "" else None
            repair_notice_model.save()
            return redirect('repair_notice')

        elif "filter" in request.POST:
            if request.POST['filter_date'] == "0":
                if request.POST['filter_status'] == "0":
                    pass
                elif request.POST['filter_status'] == "1":
                    list_repair_notice = Repair_notice.objects.filter(repairer_user=User_login).exclude(repair_status="ปิดใบแจ้งซ่อม")
                elif request.POST['filter_status'] == "2":
                    list_repair_notice = Repair_notice.objects.filter(repairer_user=User_login, repair_status="ปิดใบแจ้งซ่อม")
                elif request.POST['filter_status'] == "3":
                    list_repair_notice = Repair_notice.objects.filter(department_notifying=UserLoginDepartment)
                elif request.POST['filter_status'] == "4":
                    list_repair_notice = Repair_notice.objects.filter(department_notifying=UserLoginDepartment).exclude(repair_status="ปิดใบแจ้งซ่อม")
                elif request.POST['filter_status'] == "5":
                    list_repair_notice = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, repair_status="ปิดใบแจ้งซ่อม")
            elif request.POST['filter_date'] == "1":
                date_today = datetime.date.today()
                if request.POST['filter_status'] == "0":
                    list_repair_notice = Repair_notice.objects.filter(repairer_user=User_login, notification_date=date_today)
                elif request.POST['filter_status'] == "1":
                    list_repair_notice = Repair_notice.objects.filter(repairer_user=User_login, notification_date=date_today).exclude(repair_status="ปิดใบแจ้งซ่อม")
                elif request.POST['filter_status'] == "2":
                    list_repair_notice = Repair_notice.objects.filter(repairer_user=User_login, repair_status="ปิดใบแจ้งซ่อม", notification_date=date_today)
                elif request.POST['filter_status'] == "3":
                    list_repair_notice = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date=date_today)
                elif request.POST['filter_status'] == "4":
                    list_repair_notice = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date=date_today).exclude(repair_status="ปิดใบแจ้งซ่อม")
                elif request.POST['filter_status'] == "5":
                    list_repair_notice = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, repair_status="ปิดใบแจ้งซ่อม", notification_date=date_today)
            elif request.POST['filter_date'] == "2":
                end_date = datetime.date.today()
                start_date = datetime.date.today()-datetime.timedelta(days=7)
                if request.POST['filter_status'] == "0":
                    list_repair_notice = Repair_notice.objects.filter(repairer_user=User_login, notification_date__range=[start_date, end_date])
                elif request.POST['filter_status'] == "1":
                    list_repair_notice = Repair_notice.objects.filter(repairer_user=User_login, notification_date__range=[start_date, end_date]).exclude(repair_status="ปิดใบแจ้งซ่อม")
                elif request.POST['filter_status'] == "2":
                    list_repair_notice = Repair_notice.objects.filter(repairer_user=User_login, repair_status="ปิดใบแจ้งซ่อม", notification_date__range=[start_date, end_date])
                elif request.POST['filter_status'] == "3":
                    list_repair_notice = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date__range=[start_date, end_date])
                elif request.POST['filter_status'] == "4":
                    list_repair_notice = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date__range=[start_date, end_date]).exclude(repair_status="ปิดใบแจ้งซ่อม")
                elif request.POST['filter_status'] == "5":
                    list_repair_notice = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, repair_status="ปิดใบแจ้งซ่อม", notification_date__range=[start_date, end_date])
            elif request.POST['filter_date'] == "3":
                this_month = datetime.date.today().month
                if request.POST['filter_status'] == "0":
                    list_repair_notice = Repair_notice.objects.filter(repairer_user=User_login, notification_date__month=this_month)
                elif request.POST['filter_status'] == "1":
                    list_repair_notice = Repair_notice.objects.filter(repairer_user=User_login, notification_date__month=this_month).exclude(repair_status="ปิดใบแจ้งซ่อม")
                elif request.POST['filter_status'] == "2":
                    list_repair_notice = Repair_notice.objects.filter(repairer_user=User_login, repair_status="ปิดใบแจ้งซ่อม", notification_date__month=this_month)
                elif request.POST['filter_status'] == "3":
                    list_repair_notice = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date__month=this_month)
                elif request.POST['filter_status'] == "4":
                    list_repair_notice = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date__month=this_month).exclude(repair_status="ปิดใบแจ้งซ่อม")
                elif request.POST['filter_status'] == "5":
                    list_repair_notice = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, repair_status="ปิดใบแจ้งซ่อม", notification_date__month=this_month)
            elif request.POST['filter_date'] == "4":
                this_year = datetime.date.today().year
                if request.POST['filter_status'] == "0":
                    list_repair_notice = Repair_notice.objects.filter(repairer_user=User_login, notification_date__year=this_year)
                elif request.POST['filter_status'] == "1":
                    list_repair_notice = Repair_notice.objects.filter(repairer_user=User_login, notification_date__year=this_year).exclude(repair_status="ปิดใบแจ้งซ่อม")
                elif request.POST['filter_status'] == "2":
                    list_repair_notice = Repair_notice.objects.filter(repairer_user=User_login, repair_status="ปิดใบแจ้งซ่อม", notification_date__year=this_year)
                elif request.POST['filter_status'] == "3":
                    list_repair_notice = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date__year=this_year)
                elif request.POST['filter_status'] == "4":
                    list_repair_notice = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date__year=this_year).exclude(repair_status="ปิดใบแจ้งซ่อม")
                elif request.POST['filter_status'] == "5":
                    list_repair_notice = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, repair_status="ปิดใบแจ้งซ่อม", notification_date__year=this_year)

        elif 'export' in request.POST:
            if request.POST['department_file'] == "1":
                if request.POST['date_file'] == "1":
                    date_today = datetime.date.today()
                    if request.POST['status_file'] == "1":
                        repair_export = Repair_notice.objects.filter(repairer_user=User_login, notification_date=date_today).exclude(repair_status="ปิดใบแจ้งซ่อม")
                    elif request.POST['status_file'] == "2":
                        repair_export = Repair_notice.objects.filter(repairer_user=User_login, notification_date=date_today, repair_status="ปิดใบแจ้งซ่อม")
                    elif request.POST['status_file'] == "3":
                        repair_export = Repair_notice.objects.filter(repairer_user=User_login, notification_date=date_today)
                elif request.POST['date_file'] == "2":
                    end_date = datetime.date.today()
                    start_date = datetime.date.today()-datetime.timedelta(days=7)
                    if request.POST['status_file'] == "1":
                        repair_export = Repair_notice.objects.filter(repairer_user=User_login, notification_date__range=[start_date, end_date]).exclude(repair_status="ปิดใบแจ้งซ่อม")
                    elif request.POST['status_file'] == "2":
                        repair_export = Repair_notice.objects.filter(repairer_user=User_login, notification_date__range=[start_date, end_date], repair_status="ปิดใบแจ้งซ่อม")
                    elif request.POST['status_file'] == "3":
                        repair_export = Repair_notice.objects.filter(repairer_user=User_login, notification_date__range=[start_date, end_date])
                elif request.POST['date_file'] == "3":
                    this_month = datetime.date.today().month
                    if request.POST['status_file'] == "1":
                        repair_export = Repair_notice.objects.filter(repairer_user=User_login, notification_date__month=this_month).exclude(repair_status="ปิดใบแจ้งซ่อม")
                    elif request.POST['status_file'] == "2":
                        repair_export = Repair_notice.objects.filter(repairer_user=User_login, notification_date__month=this_month, repair_status="ปิดใบแจ้งซ่อม")
                    elif request.POST['status_file'] == "3":
                        repair_export = Repair_notice.objects.filter(repairer_user=User_login, notification_date__month=this_month)
                elif request.POST['date_file'] == "4":
                    this_year = datetime.date.today().year
                    if request.POST['status_file'] == "1":
                        repair_export = Repair_notice.objects.filter(repairer_user=User_login, notification_date__year=this_year).exclude(repair_status="ปิดใบแจ้งซ่อม")
                    elif request.POST['status_file'] == "2":
                        repair_export = Repair_notice.objects.filter(repairer_user=User_login, notification_date__year=this_year, repair_status="ปิดใบแจ้งซ่อม")
                    elif request.POST['status_file'] == "3":
                        repair_export = Repair_notice.objects.filter(repairer_user=User_login, notification_date__year=this_year)
            elif request.POST['department_file'] == "2":
                if request.POST['date_file'] == "1":
                    date_today = datetime.date.today()
                    if request.POST['status_file'] == "1":
                        repair_export = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date=date_today).exclude(repair_status="ปิดใบแจ้งซ่อม")
                    elif request.POST['status_file'] == "2":
                        repair_export = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date=date_today, repair_status="ปิดใบแจ้งซ่อม")
                    elif request.POST['status_file'] == "3":
                        repair_export = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date=date_today)
                elif request.POST['date_file'] == "2":
                    end_date = datetime.date.today()
                    start_date = datetime.date.today()-datetime.timedelta(days=7)
                    if request.POST['status_file'] == "1":
                        repair_export = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date__range=[start_date, end_date]).exclude(repair_status="ปิดใบแจ้งซ่อม")
                    elif request.POST['status_file'] == "2":
                        repair_export = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date__range=[start_date, end_date], repair_status="ปิดใบแจ้งซ่อม")
                    elif request.POST['status_file'] == "3":
                        repair_export = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date__range=[start_date, end_date])
                elif request.POST['date_file'] == "3":
                    this_month = datetime.date.today().month
                    if request.POST['status_file'] == "1":
                        repair_export = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date__month=this_month).exclude(repair_status="ปิดใบแจ้งซ่อม")
                    elif request.POST['status_file'] == "2":
                        repair_export = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date__month=this_month, repair_status="ปิดใบแจ้งซ่อม")
                    elif request.POST['status_file'] == "3":
                        repair_export = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date__month=this_month)
                elif request.POST['date_file'] == "4":
                    this_year = datetime.date.today().year
                    if request.POST['status_file'] == "1":
                        repair_export = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date__year=this_year).exclude(repair_status="ปิดใบแจ้งซ่อม")
                    elif request.POST['status_file'] == "2":
                        repair_export = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date__year=this_year, repair_status="ปิดใบแจ้งซ่อม")
                    elif request.POST['status_file'] == "3":
                        repair_export = Repair_notice.objects.filter(department_notifying=UserLoginDepartment, notification_date__year=this_year)

            document = Document()
            document.add_heading(f'ใบแจ้งซ่อมเครื่องจักร', 0)
            for index, repair in enumerate(repair_export):
                document.add_heading(f'{index+1}. หมายเลขใบแจ้งซ่อม : {repair.repair_no}', level=1)

                if repair.machine.machine_image1:
                    picture_mch = document.add_picture(repair.machine.machine_image1, width=Inches(2))
                    last_paragraph = document.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                text1 = document.add_paragraph('')
                text1.add_run('บันทึกผู้แจ้งซ่อมเครื่องจักร').bold = True
                text1.alignment = 1
                text1.underline = True
                document.add_paragraph(f'วันที่แจ้งซ่อมเครื่องจักร : {repair.notification_date}')
                document.add_paragraph(f'หน่วยงานที่แจ้งซ่อมเครื่องจักร : {repair.department_notifying.department_code} | {repair.department_notifying.department_name}')
                p = document.add_paragraph(f'สถานะใบแจ้งซ่อมเครื่องจักร : ')
                p.add_run(repair.repair_status).bold = True
                document.add_paragraph(f'ผู้แจ้งใบแจ้งซ่อมเครื่องจักร : {repair.repairer_user.firstname} {repair.repairer_user.lastname}')
                document.add_paragraph(f'ผู้ตรวจสอบใบแจ้งซ่อมเครื่องจักร : {repair.inspect_user.firstname} {repair.inspect_user.lastname}')
                document.add_paragraph(f'ผู้อนุมัติใบแจ้งซ่อมเครื่องจักร : {repair.approve_user.firstname} {repair.approve_user.lastname}')

                text2 = document.add_paragraph('')
                text2.add_run('ข้อมูลเครื่องจักรที่แจ้งซ่อม').bold = True
                text2.alignment = 1
                text2.underline = True
                document.add_paragraph(f'วันที่ต้องการใช้งาน : {repair.use_date}')
                document.add_paragraph(f'สายการผลิตที่ : {repair.machine.line}')
                document.add_paragraph(f'ชื่อเครื่องจักร : {repair.machine.machine_name} \t รหัสเครื่องจักร : {repair.machine.machine_production_line_code}')
                document.add_paragraph(f'ปัญหาเครื่องจักรที่พบ : {repair.problem_report}')
                document.add_paragraph(f'ผลกระทบของปัญหา : {repair.effect_problem}')
                if repair.maintenance_jobs:
                    text3 = document.add_paragraph()
                    text3.add_run('ข้อมูลงานการซ่อมบำรุงเครื่องจักร').bold = True
                    text3.alignment = 1

                    for number, job in enumerate(repair.maintenance_jobs.all()):
                        # document.add_heading(f'{index+1}. หมายเลขงานซ่อมบำรุง : {job.job_no}', level=2)
                        document.add_paragraph(f'{number+1}. หมายเลขงานซ่อมบำรุง : {job.job_no}')
                        document.add_paragraph(f'อะไหล่ที่ซ่อมบำรุงเครื่องจักร : {job.job_mch_sp.spare_part.spare_part_name}')
                        document.add_paragraph(f'วันที่สร้างงาน : {job.job_gen_date}')
                        document.add_paragraph(f'ประเภทของงานซ่อมบำรุง : {job.job_mtn_type}')
                        p = document.add_paragraph(f'สถานะของงานซ่อมบำรง : ')
                        p.add_run(job.job_status).bold = True
                        if job.job_response_user:
                            document.add_paragraph(f'ผู้มอบหมายงานซ่อมบำรุง : {job.job_assign_user.firstname} {job.job_assign_user.lastname}')
                            document.add_paragraph(f'ผู้รับผิดชอบงานซ่อมบำรุง : {job.job_response_user.firstname} {job.job_response_user.lastname}')
                        if job.is_approve:
                            document.add_paragraph(f'สาเหตุของปัญหา : {job.problem_cause}')
                            document.add_paragraph(f'วิธีการแก้ไขปัญหา : {job.corrective_action}')
                            document.add_paragraph(f'วิธีการดูแลหลังซ่อมบำรุง : {job.after_repair}')
                            document.add_paragraph(f'ชั่วโมงเครื่องจักรที่บันทึก : {job.job_mch_hour}')
                            document.add_paragraph(f'อายุการมาตรวจสอบ : {job.job_plan_hour}')
                            document.add_paragraph(f'อายุการมาซ่อมบำรุง : {job.job_fix_plan_hour}')
                            document.add_paragraph(f'ผลการซ่อมบำรุง : {job.job_result_type}')
                            if job.job_result_description: document.add_paragraph(f'หมายเหตุ : {job.job_result_description}')
                if index+1 != len(repair_export):
                    document.add_page_break()

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=repair_notice.docx'
            document.save(response)
            return response

    context = {'line_of_user': line_of_user, 'User_login': User_login, 'list_repair_notice': list_repair_notice,
               'list_inspect_user': list_inspect_user, 'list_approve_user': list_approve_user, "UserLoginDepartment": UserLoginDepartment,
               'list_department': list_department, 'repair_menu': repair_menu, 'repair_main_menu': repair_main_menu}
    return render(request, 'repair_inform/repair_notice.html', context)


def department_manage(request):
    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='department_manage')
    if not role_and_screen.exists():
        return redirect('signin')
    departments = Department.objects.all()
    if request.method == "POST":
        if "add_dep" in request.POST:
            create_dep = Department.objects.create(department_code=request.POST["add_dep_code"],
                                                   department_name=request.POST['add_dep_name'],
                                                   create_by=User_login.username)
            create_dep.save()
        elif "edit_dep" in request.POST:
            update_dep = Department.objects.get(department_code=request.POST['set_dep_code'])
            update_dep.department_name = request.POST['set_dep_name']
            update_dep.update_by = str(User_login.username)
            update_dep.save()
        elif "delete_dep" in request.POST:
            delete_dep = Department.objects.get(pk=request.POST['delete_dep'])
            delete_dep.delete()

        return redirect('department_manage')

    context = {'User_login': User_login, 'departments': departments}
    return render(request, 'account/department_manage.html', context)


def user_department(request):
    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='user_and_department')
    if not role_and_screen.exists():
        return redirect('signin')
    org_all = Organization.objects.all()
    user_all = User.objects.all()
    user_department_all = User_and_department.objects.all()
    department_all = Department.objects.all()
    dict_user_department = {}
    for user_dep in user_department_all:
        if user_dep.department not in dict_user_department.keys():
            dict_user_department[user_dep.department] = []
            dict_user_department[user_dep.department].append(user_dep.user)
        else:
            dict_user_department[user_dep.department].append(user_dep.user)

    if request.method == "POST":
        if 'addMembersDepartment' in request.POST:
            add_department_code = request.POST['add_department_code']
            select_user = request.POST.getlist('select_user')
            for user in select_user:
                try:
                    user_model = User.objects.get(username=user)
                    department_model = Department.objects.get(department_code=add_department_code)
                    user_model.departments.add(department_model)
                    user_model.save()
                except ObjectDoesNotExist:
                    messages.error(request, 'กรอกข้อมูลไม่ถูกต้อง กรุณาลองใหม่')

        elif 'update_user_dep' in request.POST:
            update_dep_code = request.POST['update_user_dep']
            update_add_members = request.POST.getlist('update_add_members', False)
            delete_user_list = request.POST.getlist('delete_user_list', False)
            toggle_delete = True if request.POST['toggle_delete'] == "true" else False
            if update_add_members:
                for user in update_add_members:
                    user_model = User.objects.get(username=user)
                    department_model = Department.objects.get(department_code=update_dep_code)
                    user_model.departments.add(department_model)
                    user_model.save()
            if toggle_delete:
                for user in delete_user_list:
                    user_model = User.objects.get(username=user)
                    department_model = Department.objects.get(department_code=update_dep_code)
                    user_model.departments.remove(department_model)
                    user_model.save()

        elif 'dep_permission' in request.POST:
            list_inform = request.POST.getlist('is_inform')
            list_inspect = request.POST.getlist('is_inspect')
            list_approve = request.POST.getlist('is_approve')
            list_receive = request.POST.getlist('is_receive')
            list_assign = request.POST.getlist('is_assign')
            list_report = request.POST.getlist('is_report')
            list_verify = request.POST.getlist('is_verify')
            list_close = request.POST.getlist('is_close')
            user_model = User.objects.get(pk=request.POST['dep_permission'])
            user_departments = User_and_department.objects.filter(user=user_model)
            for user_dep in user_departments:
                user_dep.is_inform = True if str(str(user_dep.department.pk)) in list_inform else False
                user_dep.is_inspect = True if str(user_dep.department.pk) in list_inspect else False
                user_dep.is_approve = True if str(user_dep.department.pk) in list_approve else False
                user_dep.is_receive = True if str(user_dep.department.pk) in list_receive else False
                user_dep.is_assign = True if str(user_dep.department.pk) in list_assign else False
                user_dep.is_report = True if str(user_dep.department.pk) in list_report else False
                user_dep.is_verify = True if str(user_dep.department.pk) in list_verify else False
                user_dep.is_close = True if str(user_dep.department.pk) in list_close else False
                user_dep.save()

        return redirect('user_department')

    context = {'User_login': User_login, 'org_all': org_all, 'user_all': user_all, 'user_department_all': user_department_all,
               'department_all': department_all, 'dict_user_department': dict_user_department}
    return render(request, 'account/user_department.html', context)


def repair_inspect(request):
    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='repair_inspect')
    if not role_and_screen.exists():
        return redirect('signin')

    repair_menu = dict_menu_level[Menu.objects.get(pk='repair_menu')]
    repair_main_menu = Menu.objects.get(pk='repair_inspect')

    repair_inspect_all = Repair_notice.objects.filter(repair_status="รอการตรวจสอบ", inspect_user=User_login)
    if request.method == "POST":
        if "repair_submit" in request.POST:
            repair_inspect_model = Repair_notice.objects.get(pk=request.POST['repair_submit'])
            if 'is_inspect' == request.POST['is_inspect']:
                repair_inspect_model.repair_status = 'รอการอนุมัติ'
                repair_inspect_model.is_inspect = True
            elif 'not_inspect' == request.POST['is_inspect']:
                repair_inspect_model.repair_status = 'ตรวจสอบไม่ผ่าน'
                repair_inspect_model.is_inspect = False
            elif 'is_cancel' == request.POST['is_inspect']:
                repair_inspect_model.repair_status = 'ยกเลิกใบแจ้ง'
                repair_inspect_model.is_inspect = False
                repair_inspect_model.is_cancel = True
            repair_inspect_model.inspect_remark = request.POST['inspect_remark'] if request.POST['inspect_remark'] != "" else None
            repair_inspect_model.save()

        return redirect('repair_inspect')

    context = {'User_login': User_login, 'repair_inspect_all': repair_inspect_all, 'repair_menu': repair_menu, 'repair_main_menu': repair_main_menu}
    return render(request, 'repair_inform/repair_inspect.html', context)


def repair_approve(request):
    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='repair_approve')
    if not role_and_screen.exists():
        return redirect('signin')

    repair_menu = dict_menu_level[Menu.objects.get(pk='repair_menu')]
    repair_main_menu = Menu.objects.get(pk='repair_approve')

    repair_approve_all = Repair_notice.objects.filter(repair_status="รอการอนุมัติ", approve_user=User_login)
    if request.method == "POST":
        if "repair_submit" in request.POST:
            repair_approve_model = Repair_notice.objects.get(pk=request.POST['repair_submit'])
            if 'is_approve' == request.POST['is_approve']:
                repair_approve_model.repair_status = 'รอการรับใบแจ้ง'
                repair_approve_model.is_approve = True
            elif 'not_approve' == request.POST['is_approve']:
                repair_approve_model.repair_status = 'อนุมัติไม่ผ่าน'
                repair_approve_model.is_approve = False
            elif 'is_cancel' == request.POST['is_approve']:
                repair_approve_model.repair_status = 'ยกเลิกใบแจ้ง'
                repair_approve_model.is_approve = False
                repair_approve_model.is_cancel = True
            repair_approve_model.approve_remark = request.POST['approve_remark'] if request.POST['approve_remark'] != "" else None
            repair_approve_model.save()

        return redirect('repair_approve')

    context = {'User_login': User_login, 'repair_approve_all': repair_approve_all, 'repair_menu': repair_menu, 'repair_main_menu': repair_main_menu}
    return render(request, 'repair_inform/repair_approve.html', context)


def signIn_department(request):
    global UserLoginDepartment
    if User_login is None:
        return redirect('signin')
    departments_user = User_login.departments.all()
    context = {"departments_user": departments_user}
    if request.method == "POST":
        if 'department_submit' in request.POST:
            UserLoginDepartment = Department.objects.get(pk=request.POST['select_department'])
            return redirect('home')

    return render(request, 'home/signIn_department.html', context)


def maintenance_receive(request):
    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='maintenance_receive')
    if not role_and_screen.exists():
        return redirect('signin')

    mtn_menu = dict_menu_level[Menu.objects.get(pk='preventive_data')]
    mtn_main_menu = Menu.objects.get(pk='maintenance_receive')

    # repair_receive_all = Repair_notice.objects.filter(repair_status="รอการรับใบแจ้ง", department_receive=UserLoginDepartment)
    repair_receive_all = Repair_notice.objects.filter(department_receive=UserLoginDepartment)
    if request.method == "POST":
        if "repair_submit" in request.POST:
            repair_receive_model = Repair_notice.objects.get(pk=request.POST['repair_submit'])
            if 'is_receive' == request.POST['is_receive']:
                repair_receive_model.repair_status = 'รอการตรวจสอบอะไหล่'
                repair_receive_model.is_receive = True
            elif 'not_receive' == request.POST['is_receive']:
                repair_receive_model.repair_status = 'ไม่ผ่านการรับแจ้ง'
                repair_receive_model.is_receive = False
            elif 'is_cancel' == request.POST['is_receive']:
                repair_receive_model.repair_status = 'ยกเลิกใบแจ้ง'
                repair_receive_model.is_receive = False
                repair_receive_model.is_cancel = True
            repair_receive_model.receive_remark = request.POST['receive_remark'] if request.POST['receive_remark'] != "" else None
            repair_receive_model.save()

        return redirect('maintenance_receive')

    context = {'User_login': User_login, 'repair_receive_all': repair_receive_all, 'mtn_menu': mtn_menu, 'mtn_main_menu': mtn_main_menu}
    return render(request, 'maintenance/maintenance_receive.html', context)


def maintenance_inspect(request):
    role_and_screen = Role_Screen.objects.filter(role_id=UserRole, screen_id='maintenance_inspect')
    if not role_and_screen.exists():
        return redirect('signin')

    mtn_menu = dict_menu_level[Menu.objects.get(pk='preventive_data')]
    mtn_main_menu = Menu.objects.get(pk='maintenance_inspect')

    # mtn_inspect_all = Repair_notice.objects.filter(repair_status="รอการตรวจสอบอะไหล่", department_receive=UserLoginDepartment)
    mtn_inspect_all = Repair_notice.objects.filter(department_receive=UserLoginDepartment)
    spare_part_group_all = Spare_part_group.objects.all()

    if request.method == "POST":
        if "mtn_submit" in request.POST:
            repair_notice_model = Repair_notice.objects.get(pk=request.POST['mtn_submit'])
            repair_notice_model.repair_status = 'อยู่ในระหว่างการทำงาน'

            list_check = request.POST.getlist('list_spare_part')
            if len(list_check) != len(set(list_check)):
                messages.error(request, "ทำรายการไม่สำเร็จ เนื่องจากเลือกชิ้นส่วนอะไหล่ซ้ำ")
                return redirect('maintenance_inspect')
            elif "0" in list_check:
                messages.error(request, "ทำรายการไม่สำเร็จ เนื่องจากไม่ได้เลือกชิ้นส่วนอะไหล่")
                return redirect('maintenance_inspect')
            for sp_id_test in list_check:
                try:
                    mch_sp = Machine_sparepart.objects.get(machine=repair_notice_model.machine, spare_part_id=sp_id_test)
                except ObjectDoesNotExist:
                    messages.error(request, "ทำรายการไม่สำเร็จ เนื่องจากเครื่องจักรยังไม่ได้เชื่อมต่อกับอะไหล่")
                    return redirect('maintenance_inspect')
                if mch_sp.gen_mtnchng_date is not None:
                    job_auto = Maintenance_job.objects.exclude(job_status__in=["ปิดงาน", "งานเสร็จสิ้น"]).get(job_mch_sp=mch_sp, job_gen_date=mch_sp.gen_mtnchng_date)
                    job_auto.job_status = "ปิดงาน"
                    job_auto.job_remark = "เนื่องจากมีการสร้างงานนี้ในใบแจ้งซ่อม"
                    job_auto.is_report = True
                    job_auto.is_approve = True
                    job_auto.save()
                elif mch_sp.gen_mtnchk_date is not None:
                    job_auto = Maintenance_job.objects.exclude(job_status__in=["ปิดงาน", "งานเสร็จสิ้น"]).get(job_mch_sp=mch_sp, job_gen_date=mch_sp.gen_mtnchk_date)
                    job_auto.job_status = "ปิดงาน"
                    job_auto.job_remark = "เนื่องจากมีการสร้างงานนี้ในใบแจ้งซ่อม"
                    job_auto.is_report = True
                    job_auto.is_approve = True
                    job_auto.save()
                mch_sp.gen_mtnchng_date = datetime.date.today()
                mch_sp.save()
                mtn_repair_job = Maintenance_job.objects.create(job_no=genJobNumber(),
                                                                job_mtn_type="repair",
                                                                job_gen_user=User_login,
                                                                job_status="รอการมอบหมายงาน",
                                                                job_mch_sp=mch_sp)
                repair_notice_model.maintenance_jobs.add(mtn_repair_job)
                mtn_repair_job.save()
            repair_notice_model.save()
            messages.success(request, "บันทึกรายการสำเร็จ")
            return redirect('maintenance_inspect')

    context = {'User_login': User_login, 'mtn_inspect_all': mtn_inspect_all, 'spare_part_group_all': spare_part_group_all,
               'mtn_menu': mtn_menu, 'mtn_main_menu': mtn_main_menu}
    return render(request, 'maintenance/maintenance_inspect.html', context)

